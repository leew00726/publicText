import io
import unittest
from unittest.mock import MagicMock, patch

from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt

from app.services.topic_inference import extract_docx_features, extract_pdf_features, infer_topic_rules


def _build_docx_bytes(body_font: str, heading_font: str, heading_size_pt: float, body_size_pt: float) -> bytes:
    doc = Document()

    heading = doc.add_paragraph("一、总体要求")
    heading.style = doc.styles["Heading 1"]
    heading_run = heading.runs[0]
    heading_run.font.name = heading_font
    heading_run.font.size = Pt(heading_size_pt)

    body = doc.add_paragraph("这是正文第一段。")
    body_run = body.runs[0]
    body_run.font.name = body_font
    body_run.font.size = Pt(body_size_pt)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.line_spacing = Pt(28)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _build_docx_with_fixed_blocks() -> bytes:
    doc = Document()
    doc.add_paragraph("华能云成数字产融科技（雄安）有限公司")
    doc.add_paragraph("周例会会议纪要")
    doc.add_paragraph("2026年第6期")
    doc.add_paragraph("综合管理部  2026年2月9日  签发人：")
    doc.add_paragraph("2月9日，公司组织召开第六期周例会，具体纪要如下：")
    doc.add_paragraph("一、战略推进与业务发展")
    doc.add_paragraph("主 持：金刚善")
    doc.add_paragraph("参 加：江进、何宣")
    doc.add_paragraph("记 录：靳冰洁")
    doc.add_paragraph("发送：全体员工。")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _build_docx_with_styled_fixed_blocks() -> bytes:
    doc = Document()

    p1 = doc.add_paragraph("华能云成数字产融科技（雄安）有限公司")
    p1.alignment = 1
    r1 = p1.runs[0]
    r1.font.name = "方正小标宋简体"
    r1.font.size = Pt(22)
    r1.font.color.rgb = RGBColor.from_string("FF0000")

    p2 = doc.add_paragraph("周例会会议纪要")
    p2.alignment = 1
    r2 = p2.runs[0]
    r2.font.name = "方正小标宋简体"
    r2.font.size = Pt(26)
    r2.font.color.rgb = RGBColor.from_string("FF0000")

    p3 = doc.add_paragraph("2026年第6期")
    p3.alignment = 1
    r3 = p3.runs[0]
    r3.font.name = "黑体"
    r3.font.size = Pt(16)

    p4 = doc.add_paragraph("2月9日，公司组织召开第六期周例会，具体纪要如下：")
    r4 = p4.runs[0]
    r4.font.name = "仿宋_GB2312"
    r4.font.size = Pt(16)
    p4.paragraph_format.first_line_indent = Pt(32)

    p5 = doc.add_paragraph("一、战略推进与业务发展")
    r5 = p5.runs[0]
    r5.font.name = "黑体"
    r5.font.size = Pt(16)
    p5.paragraph_format.first_line_indent = Pt(32)

    p6 = doc.add_paragraph("发送：全体员工。")
    r6 = p6.runs[0]
    r6.font.name = "仿宋_GB2312"
    r6.font.size = Pt(16)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _build_docx_with_mixed_suffix_fonts() -> bytes:
    doc = Document()

    title = doc.add_paragraph("周例会会议纪要")
    title.runs[0].font.name = "方正小标宋简体"
    title.runs[0].font.size = Pt(22)

    body = doc.add_paragraph("2月9日，公司组织召开第六期周例会。")
    body.runs[0].font.name = "仿宋_GB2312"
    body.runs[0].font.size = Pt(16)
    body.paragraph_format.first_line_indent = Pt(32)
    body.paragraph_format.line_spacing = Pt(28)

    body2 = doc.add_paragraph("会议强调要按计划推进重点项目。")
    body2.runs[0].font.name = "仿宋_GB2312"
    body2.runs[0].font.size = Pt(16)
    body2.paragraph_format.first_line_indent = Pt(32)
    body2.paragraph_format.line_spacing = Pt(28)

    host = doc.add_paragraph("主 持：金刚善")
    host.runs[0].font.name = "黑体"
    host.runs[0].font.size = Pt(16)

    attendees = doc.add_paragraph("参会人员：张三、李四")
    attendees.runs[0].font.name = "黑体"
    attendees.runs[0].font.size = Pt(16)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _build_docx_without_body_with_header_and_suffix() -> bytes:
    doc = Document()
    doc.add_paragraph("普通商密★1年")
    doc.add_paragraph("华能云成数字产融科技（雄安）有限公司")
    doc.add_paragraph("总经理办公会会议纪要")
    doc.add_paragraph("2026年第1期")
    doc.add_paragraph("综合管理部  2026年1月12日  签发人：")
    doc.add_paragraph("主  持：汪  进")
    doc.add_paragraph("参  加：何  亘、段国强、杨立寨")
    doc.add_paragraph("记  录：李  健")
    doc.add_paragraph("发送：董事长，总经理办公会成员。")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


class TopicInferenceTests(unittest.TestCase):
    def test_extract_docx_features_reads_basic_styles(self) -> None:
        content = _build_docx_bytes("仿宋_GB2312", "黑体", 16, 16)
        features = extract_docx_features(content)

        self.assertEqual(features["body"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(features["headings"]["level1"]["fontFamily"], "黑体")
        self.assertGreater(features["page"]["marginsCm"]["top"], 0)

    def test_infer_topic_rules_computes_mode_and_confidence(self) -> None:
        f1 = extract_docx_features(_build_docx_bytes("仿宋_GB2312", "黑体", 16, 16))
        f2 = extract_docx_features(_build_docx_bytes("仿宋_GB2312", "黑体", 16, 16))
        f3 = extract_docx_features(_build_docx_bytes("宋体", "黑体", 16, 16))

        rules, confidence = infer_topic_rules([f1, f2, f3])

        self.assertEqual(rules["body"]["fontFamily"], "仿宋_GB2312")
        self.assertAlmostEqual(confidence["body.fontFamily"]["confidence"], 2 / 3, places=2)
        self.assertEqual(rules["headings"]["level1"]["fontFamily"], "黑体")

    def test_infer_topic_rules_rejects_empty_features(self) -> None:
        with self.assertRaises(ValueError):
            infer_topic_rules([])

    def test_extract_docx_features_includes_content_template_blocks(self) -> None:
        features = extract_docx_features(_build_docx_with_fixed_blocks())

        content_template = features.get("contentTemplate")
        self.assertIsInstance(content_template, dict)
        self.assertGreaterEqual(len(content_template.get("leadingNodes", [])), 3)
        leading_texts = [
            "".join(part.get("text", "") for part in (node.get("content") or []))
            for node in content_template.get("leadingNodes", [])
            if isinstance(node, dict)
        ]
        self.assertTrue(any("签发人" in text for text in leading_texts))
        trailing_texts = [
            "".join(part.get("text", "") for part in (node.get("content") or []))
            for node in content_template.get("trailingNodes", [])
            if isinstance(node, dict)
        ]
        self.assertTrue(any(text.startswith("主 持") for text in trailing_texts))
        self.assertTrue(any(text.startswith("发送") for text in trailing_texts))

    def test_extract_docx_features_keeps_fixed_block_styles_and_excludes_from_body_mode(self) -> None:
        features = extract_docx_features(_build_docx_with_styled_fixed_blocks())

        self.assertEqual(features["body"]["fontFamily"], "仿宋_GB2312")

        content_template = features.get("contentTemplate") or {}
        leading = content_template.get("leadingNodes") or []
        self.assertGreaterEqual(len(leading), 3)

        first_attrs = leading[0].get("attrs") or {}
        self.assertEqual(first_attrs.get("textAlign"), "center")
        self.assertEqual(first_attrs.get("colorHex"), "#FF0000")
        self.assertEqual(first_attrs.get("fontFamily"), "方正小标宋简体")

    def test_infer_topic_rules_normalizes_suffix_line_font_to_body(self) -> None:
        features = extract_docx_features(_build_docx_with_mixed_suffix_fonts())
        rules, _ = infer_topic_rules([features])

        body_font = (rules.get("body") or {}).get("fontFamily")
        self.assertEqual(body_font, "仿宋_GB2312")

        content_template = rules.get("contentTemplate") or {}
        trailing = content_template.get("trailingNodes") or []
        host_node = next(
            (
                node
                for node in trailing
                if isinstance(node, dict)
                and "".join(part.get("text", "") for part in (node.get("content") or []) if isinstance(part, dict)).startswith("主 持")
            ),
            None,
        )
        self.assertIsNotNone(host_node)
        self.assertEqual((host_node.get("attrs") or {}).get("fontFamily"), "仿宋_GB2312")

        attendee_node = next(
            (
                node
                for node in trailing
                if isinstance(node, dict)
                and "".join(part.get("text", "") for part in (node.get("content") or []) if isinstance(part, dict)).startswith("参会人员")
            ),
            None,
        )
        self.assertIsNotNone(attendee_node)
        self.assertEqual((attendee_node.get("attrs") or {}).get("fontFamily"), "仿宋_GB2312")

    def test_extract_docx_features_fallback_keeps_leading_nodes_when_body_missing(self) -> None:
        features = extract_docx_features(_build_docx_without_body_with_header_and_suffix())
        content_template = features.get("contentTemplate") or {}
        leading = content_template.get("leadingNodes") or []
        trailing = content_template.get("trailingNodes") or []

        self.assertGreaterEqual(len(leading), 5)
        leading_texts = [
            "".join(part.get("text", "") for part in (node.get("content") or []))
            for node in leading
            if isinstance(node, dict)
        ]
        self.assertTrue(any("华能云成数字产融科技" in text for text in leading_texts))
        self.assertTrue(any("签发人" in text for text in leading_texts))
        self.assertTrue(any((node.get("attrs") or {}).get("dividerRed") for node in leading if isinstance(node, dict)))

        trailing_texts = [
            "".join(part.get("text", "") for part in (node.get("content") or []))
            for node in trailing
            if isinstance(node, dict)
        ]
        self.assertTrue(any(text.startswith("主") for text in trailing_texts))
        trailing_dividers = [
            node for node in trailing if isinstance(node, dict) and (node.get("attrs") or {}).get("dividerRed")
        ]
        self.assertGreaterEqual(len(trailing_dividers), 2)

    @patch("app.services.topic_inference.PdfReader")
    def test_extract_pdf_features_reads_font_and_size(self, mock_reader) -> None:
        page = MagicMock()

        def fake_extract_text(visitor_text=None):
            if visitor_text:
                visitor_text("标题", None, None, {"/BaseFont": "/ABCDEE+SimHei"}, 16)
                visitor_text("正文", None, None, {"/BaseFont": "/ABCDEE+FangSong_GB2312"}, 14)
                visitor_text("正文2", None, None, {"/BaseFont": "/ABCDEE+FangSong_GB2312"}, 14)
            return "标题 正文 正文2"

        page.extract_text.side_effect = fake_extract_text
        mock_reader.return_value.pages = [page]

        features = extract_pdf_features(b"%PDF-1.4")
        self.assertEqual(features["body"]["fontFamily"], "FangSong_GB2312")
        self.assertEqual(features["body"]["fontSizePt"], 14)

    @patch("app.services.topic_inference.PdfReader")
    def test_extract_pdf_features_rejects_scanned_pdf(self, mock_reader) -> None:
        page = MagicMock()
        page.extract_text.return_value = ""
        mock_reader.return_value.pages = [page]

        with self.assertRaises(ValueError):
            extract_pdf_features(b"%PDF-1.4")


if __name__ == "__main__":
    unittest.main()
