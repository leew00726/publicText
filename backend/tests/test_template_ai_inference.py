import io
import unittest
from unittest.mock import patch

from docx import Document
from docx.shared import Pt

from app.config import Settings
from app.services.ai_agent import AgentUpstreamError
from app.services.template_ai_inference import infer_topic_rules_with_optional_deepseek
from app.services.topic_inference import extract_docx_features


def _build_role_sample_docx() -> bytes:
    doc = Document()

    secret = doc.add_paragraph("普通商密★1年")
    secret.runs[0].font.name = "黑体"
    secret.runs[0].font.size = Pt(16)

    title = doc.add_paragraph("总经理办公会会议纪要")
    title.alignment = 1
    title.runs[0].font.name = "方正小标宋简体"
    title.runs[0].font.size = Pt(26)

    issue = doc.add_paragraph("2026年第1期")
    issue.alignment = 1
    issue.runs[0].font.name = "黑体"
    issue.runs[0].font.size = Pt(16)

    intro = doc.add_paragraph("2026年1月12日，公司召开会议。现将会议议定事项纪要如下：")
    intro.runs[0].font.name = "仿宋_GB2312"
    intro.runs[0].font.size = Pt(16)

    topic = doc.add_paragraph("专题A：请示")
    topic.runs[0].font.name = "黑体"
    topic.runs[0].font.size = Pt(16)

    body = doc.add_paragraph("同意《》。")
    body.runs[0].font.name = "仿宋_GB2312"
    body.runs[0].font.size = Pt(16)

    suffix = doc.add_paragraph("主  持：汪  进")
    suffix.runs[0].font.name = "黑体"
    suffix.runs[0].font.size = Pt(16)

    payload = io.BytesIO()
    doc.save(payload)
    return payload.getvalue()


def _build_simple_docx() -> bytes:
    doc = Document()
    heading = doc.add_paragraph("一、总体要求")
    heading.runs[0].font.name = "黑体"
    body = doc.add_paragraph("这是正文第一段。")
    body.runs[0].font.name = "仿宋_GB2312"

    payload = io.BytesIO()
    doc.save(payload)
    return payload.getvalue()


def _build_semantic_role_docx() -> bytes:
    doc = Document()

    title = doc.add_paragraph("关于供应链金融业务方案的请示")
    title.alignment = 1
    title.runs[0].font.name = "方正小标宋简体"
    title.runs[0].font.size = Pt(22)

    recipient = doc.add_paragraph("主送：董事会")
    recipient.runs[0].font.name = "仿宋_GB2312"
    recipient.runs[0].font.size = Pt(16)

    body = doc.add_paragraph("现将有关事项请示如下。")
    body.runs[0].font.name = "仿宋_GB2312"
    body.runs[0].font.size = Pt(16)

    attachment = doc.add_paragraph("附件：供应链金融业务实施方案")
    attachment.runs[0].font.name = "仿宋_GB2312"
    attachment.runs[0].font.size = Pt(16)

    signature = doc.add_paragraph("贵诚信托")
    signature.alignment = 2
    signature.runs[0].font.name = "仿宋_GB2312"
    signature.runs[0].font.size = Pt(16)

    reference = doc.add_paragraph("引用：《业务管理办法》（贵诚发〔2026〕1号）")
    reference.runs[0].font.name = "仿宋_GB2312"
    reference.runs[0].font.size = Pt(16)

    payload = io.BytesIO()
    doc.save(payload)
    return payload.getvalue()


class TemplateAiInferenceTests(unittest.TestCase):
    def test_rules_mode_does_not_call_deepseek(self) -> None:
        features = extract_docx_features(_build_simple_docx())
        settings = Settings(template_inference_engine="rules")

        with patch("app.services.template_ai_inference.classify_template_layout_with_deepseek") as classify:
            rules, confidence = infer_topic_rules_with_optional_deepseek([features], settings=settings)

        classify.assert_not_called()
        self.assertEqual(confidence["templateInference"]["engine"], "rules")
        self.assertEqual(rules["body"]["fontFamily"], "仿宋_GB2312")

    def test_hybrid_mode_uses_deepseek_roles_but_keeps_local_styles(self) -> None:
        features = extract_docx_features(_build_role_sample_docx())
        settings = Settings(template_inference_engine="hybrid", deepseek_api_key="test-key")
        ai_roles = {
            "summary": "识别会议纪要主标题、正文、专题题头和尾部名单。",
            "files": [
                {
                    "fileIndex": 0,
                    "titleIndex": 1,
                    "bodyIndexes": [3, 5],
                    "headings": [{"index": 4, "level": 1}],
                    "leadingIndexes": [0, 2],
                    "trailingIndexes": [6],
                }
            ],
        }

        with patch(
            "app.services.template_ai_inference.classify_template_layout_with_deepseek",
            return_value=ai_roles,
        ):
            rules, confidence = infer_topic_rules_with_optional_deepseek([features], settings=settings)

        self.assertEqual(confidence["templateInference"]["engine"], "deepseek")
        self.assertEqual(confidence["templateInference"]["summary"], ai_roles["summary"])
        self.assertEqual(rules["title"]["fontFamily"], "方正小标宋简体")
        self.assertEqual(rules["title"]["fontSizePt"], 26.0)
        self.assertEqual(rules["body"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(rules["headings"]["level1"]["fontFamily"], "黑体")
        self.assertNotEqual(rules["headings"]["level1"].get("fontFamily"), "模型编造字体")

    def test_deepseek_roles_preserve_semantic_sections_with_local_nodes(self) -> None:
        features = extract_docx_features(_build_semantic_role_docx())
        settings = Settings(template_inference_engine="hybrid", deepseek_api_key="test-key")
        ai_roles = {
            "summary": "识别请示模板的主送、附件、落款和引用规则。",
            "files": [
                {
                    "fileIndex": 0,
                    "titleIndex": 0,
                    "bodyIndexes": [2],
                    "recipientsIndexes": [1],
                    "attachmentIndexes": [3],
                    "signatureIndexes": [4],
                    "referenceIndexes": [5],
                }
            ],
        }

        with patch(
            "app.services.template_ai_inference.classify_template_layout_with_deepseek",
            return_value=ai_roles,
        ):
            rules, confidence = infer_topic_rules_with_optional_deepseek([features], settings=settings)

        self.assertEqual(rules["body"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(rules["title"]["fontFamily"], "方正小标宋简体")
        self.assertEqual(confidence["templateInference"]["roleSummary"]["recipients"], 1)

        content_roles = rules["contentRoles"]
        self.assertEqual(content_roles["recipients"][0]["text"], "主送：董事会")
        self.assertEqual(content_roles["attachments"][0]["text"], "附件：供应链金融业务实施方案")
        self.assertEqual(content_roles["signature"][0]["text"], "贵诚信托")
        self.assertIn("贵诚发〔2026〕1号", content_roles["references"][0]["text"])
        self.assertEqual(content_roles["signature"][0]["node"]["attrs"]["textAlign"], "right")
        self.assertNotIn("模型编造字体", str(content_roles))

    def test_hybrid_mode_falls_back_to_rules_when_deepseek_fails(self) -> None:
        features = extract_docx_features(_build_simple_docx())
        settings = Settings(template_inference_engine="hybrid", deepseek_api_key="test-key")

        with patch(
            "app.services.template_ai_inference.classify_template_layout_with_deepseek",
            side_effect=AgentUpstreamError("bad json"),
        ):
            rules, confidence = infer_topic_rules_with_optional_deepseek([features], settings=settings)

        self.assertEqual(confidence["templateInference"]["engine"], "rules")
        self.assertEqual(confidence["templateInference"]["aiStatus"], "fallback")
        self.assertIn("bad json", confidence["templateInference"]["fallbackReason"])
        self.assertEqual(rules["body"]["fontFamily"], "仿宋_GB2312")


if __name__ == "__main__":
    unittest.main()
