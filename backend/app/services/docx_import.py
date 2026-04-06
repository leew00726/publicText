from __future__ import annotations

import io
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.checker import normalize_doc_no_brackets

RE_H1 = re.compile(r"^[一二三四五六七八九十百千]+、")
RE_H2 = re.compile(r"^（[一二三四五六七八九十百千]+）")
RE_H3 = re.compile(r"^\d+\.")
RE_H4 = re.compile(r"^（\d+）")
RE_SENTENCE_PUNCT = re.compile(r"[。！？；：:]$")

ALIGNMENT_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
    WD_ALIGN_PARAGRAPH.JUSTIFY_MED: "justify",
    WD_ALIGN_PARAGRAPH.JUSTIFY_HI: "justify",
    WD_ALIGN_PARAGRAPH.JUSTIFY_LOW: "justify",
    WD_ALIGN_PARAGRAPH.THAI_JUSTIFY: "justify",
}


def _visible_runs(paragraph) -> list[Any]:
    return [run for run in paragraph.runs if (run.text or "").strip()]


def _iter_style_fonts(paragraph):
    style = getattr(paragraph, "style", None)
    while style is not None:
        font = getattr(style, "font", None)
        if font is not None:
            yield font
        style = getattr(style, "base_style", None)


def _iter_style_paragraph_formats(paragraph):
    style = getattr(paragraph, "style", None)
    while style is not None:
        paragraph_format = getattr(style, "paragraph_format", None)
        if paragraph_format is not None:
            yield paragraph_format
        style = getattr(style, "base_style", None)


def _font_name(paragraph) -> str:
    visible_runs = _visible_runs(paragraph)
    if visible_runs:
        weighted: dict[str, int] = {}
        for run in visible_runs:
            font_name = (run.font.name or "").strip()
            if not font_name:
                continue
            weighted[font_name] = weighted.get(font_name, 0) + len((run.text or "").strip())
        if weighted:
            return max(weighted.items(), key=lambda item: item[1])[0]

    for run in paragraph.runs:
        if run.font and run.font.name:
            return run.font.name
    for font in _iter_style_fonts(paragraph):
        if font.name:
            return str(font.name).strip()
    return ""


def _font_size_pt(paragraph) -> float | None:
    visible_runs = _visible_runs(paragraph)
    if visible_runs:
        weighted: dict[float, int] = {}
        for run in visible_runs:
            if not run.font or not run.font.size:
                continue
            size_pt = float(run.font.size.pt)
            weighted[size_pt] = weighted.get(size_pt, 0) + len((run.text or "").strip())
        if weighted:
            return max(weighted.items(), key=lambda item: item[1])[0]

    for run in paragraph.runs:
        if run.font and run.font.size:
            return float(run.font.size.pt)
    for font in _iter_style_fonts(paragraph):
        if font.size:
            return float(font.size.pt)
    return None


def _font_color_hex(paragraph) -> str | None:
    visible_runs = _visible_runs(paragraph)
    if visible_runs:
        weighted: dict[str, int] = {}
        for run in visible_runs:
            color = getattr(getattr(run.font, "color", None), "rgb", None)
            if color is None:
                continue
            weighted[str(color)] = weighted.get(str(color), 0) + len((run.text or "").strip())
        if weighted:
            return f"#{max(weighted.items(), key=lambda item: item[1])[0]}"

    for run in paragraph.runs:
        color = getattr(getattr(run.font, "color", None), "rgb", None)
        if color is not None:
            return f"#{color}"
    for font in _iter_style_fonts(paragraph):
        color = getattr(getattr(font, "color", None), "rgb", None)
        if color is not None:
            return f"#{color}"
    return None


def _paragraph_bold(paragraph) -> bool | None:
    visible_runs = _visible_runs(paragraph)
    if visible_runs:
        weighted: dict[bool, int] = {}
        for run in visible_runs:
            value = run.bold
            if value is None:
                value = getattr(run.font, "bold", None)
            if value is None:
                continue
            weighted[bool(value)] = weighted.get(bool(value), 0) + len((run.text or "").strip())
        if weighted:
            return max(weighted.items(), key=lambda item: item[1])[0]

    for run in paragraph.runs:
        value = run.bold
        if value is None:
            value = getattr(run.font, "bold", None)
        if value is not None:
            return bool(value)
    for font in _iter_style_fonts(paragraph):
        value = getattr(font, "bold", None)
        if value is not None:
            return bool(value)
    return None


def _length_pt(value: Any) -> float | None:
    if value is None:
        return None
    pt = getattr(value, "pt", None)
    if pt is not None:
        return round(float(pt), 2)
    return None


def _paragraph_alignment_name(paragraph) -> str | None:
    alignment = paragraph.alignment
    if alignment is None:
        for paragraph_format in _iter_style_paragraph_formats(paragraph):
            alignment = paragraph_format.alignment
            if alignment is not None:
                break
    return ALIGNMENT_MAP.get(alignment)


def _paragraph_metric_pt(paragraph, attr_name: str) -> float | None:
    direct_value = getattr(paragraph.paragraph_format, attr_name)
    direct_pt = _length_pt(direct_value)
    if direct_pt is not None:
        return direct_pt

    for paragraph_format in _iter_style_paragraph_formats(paragraph):
        style_pt = _length_pt(getattr(paragraph_format, attr_name))
        if style_pt is not None:
            return style_pt
    return None


def _paragraph_line_spacing_pt(paragraph, font_size_pt: float | None) -> float | None:
    direct_value = paragraph.paragraph_format.line_spacing
    direct_pt = _length_pt(direct_value)
    if direct_pt is not None:
        return direct_pt
    if isinstance(direct_value, (int, float)) and font_size_pt:
        return round(float(direct_value) * font_size_pt, 2)

    for paragraph_format in _iter_style_paragraph_formats(paragraph):
        style_value = paragraph_format.line_spacing
        style_pt = _length_pt(style_value)
        if style_pt is not None:
            return style_pt
        if isinstance(style_value, (int, float)) and font_size_pt:
            return round(float(style_value) * font_size_pt, 2)
    return None


def _extract_paragraph_attrs(paragraph, *, default_indent_chars: float | None = None) -> dict[str, Any]:
    attrs: dict[str, Any] = {}

    align = _paragraph_alignment_name(paragraph)
    if align:
        attrs["textAlign"] = align

    font_name = _font_name(paragraph)
    if font_name:
        attrs["fontFamily"] = font_name

    font_size_pt = _font_size_pt(paragraph)
    if font_size_pt is not None:
        attrs["fontSizePt"] = round(font_size_pt, 2)

    bold = _paragraph_bold(paragraph)
    if bold is not None:
        attrs["bold"] = bold

    color_hex = _font_color_hex(paragraph)
    if color_hex:
        attrs["colorHex"] = color_hex

    line_spacing_pt = _paragraph_line_spacing_pt(paragraph, font_size_pt)
    if line_spacing_pt is not None:
        attrs["lineSpacingPt"] = line_spacing_pt

    first_line_indent_pt = _paragraph_metric_pt(paragraph, "first_line_indent")
    if first_line_indent_pt is not None:
        attrs["firstLineIndentPt"] = first_line_indent_pt
    elif default_indent_chars is not None:
        attrs["firstLineIndentChars"] = default_indent_chars
    else:
        attrs["firstLineIndentPt"] = 0

    left_indent_pt = _paragraph_metric_pt(paragraph, "left_indent")
    if left_indent_pt is not None:
        attrs["leftIndentPt"] = left_indent_pt

    right_indent_pt = _paragraph_metric_pt(paragraph, "right_indent")
    if right_indent_pt is not None:
        attrs["rightIndentPt"] = right_indent_pt

    space_before_pt = _paragraph_metric_pt(paragraph, "space_before")
    if space_before_pt is not None:
        attrs["spaceBeforePt"] = space_before_pt

    space_after_pt = _paragraph_metric_pt(paragraph, "space_after")
    if space_after_pt is not None:
        attrs["spaceAfterPt"] = space_after_pt

    return attrs


def _looks_like_heading_candidate(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return False
    if not 2 <= len(compact) <= 40:
        return False
    if RE_SENTENCE_PUNCT.search(compact):
        return False
    if compact.endswith("："):
        return False
    if len(re.findall(r"[，,；;]", compact)) > 1:
        return False
    return True


def _looks_like_main_title(text: str, paragraph) -> bool:
    compact = re.sub(r"\s+", "", text)
    if not compact:
        return False
    if not 6 <= len(compact) <= 48:
        return False
    if RE_SENTENCE_PUNCT.search(compact):
        return False

    font_name = _font_name(paragraph)
    font_size = _font_size_pt(paragraph) or 0
    is_centered = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER

    return is_centered and ("小标宋" in font_name or "标宋" in font_name or font_size >= 18)


def _detect_level(text: str, paragraph) -> int | None:
    if RE_H1.match(text):
        return 1
    if RE_H2.match(text):
        return 2
    if RE_H3.match(text):
        return 3
    if RE_H4.match(text):
        return 4

    font_name = _font_name(paragraph)
    if "黑体" in font_name and _looks_like_heading_candidate(text):
        return 1
    if "楷体" in font_name and _looks_like_heading_candidate(text):
        return 2
    if "仿宋" in font_name and _looks_like_heading_candidate(text):
        return 3
    return None


def _make_text_node(text: str) -> dict[str, Any]:
    return {"type": "text", "text": text}


def _make_heading(level: int, text: str, attrs: dict[str, Any] | None = None) -> dict[str, Any]:
    merged_attrs = {"level": level, **(attrs or {})}
    return {
        "type": "heading",
        "attrs": merged_attrs,
        "content": [_make_text_node(text)],
    }


def _make_paragraph(text: str, attrs: dict[str, Any] | None = None) -> dict[str, Any]:
    return {
        "type": "paragraph",
        "attrs": attrs or {},
        "content": [_make_text_node(text)],
    }


def _table_to_node(table) -> dict[str, Any]:
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            cells.append({"type": "tableCell", "content": [_make_text_node(cell_text)]})
        rows.append({"type": "tableRow", "content": cells})
    return {"type": "table", "content": rows}


def _parse_numbering_warning(headings: list[tuple[int, str]]) -> list[str]:
    warnings: list[str] = []

    expected = {1: 0, 2: 0, 3: 0, 4: 0}

    def zh_to_num(zh: str) -> int:
        m = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9, "十": 10}
        if zh == "十":
            return 10
        if zh.startswith("十"):
            return 10 + m.get(zh[1:], 0)
        if "十" in zh:
            left, _, right = zh.partition("十")
            return m.get(left, 0) * 10 + m.get(right, 0)
        return m.get(zh, 0)

    for i, (level, text) in enumerate(headings):
        for lv in range(level + 1, 5):
            expected[lv] = 0
        expected[level] += 1

        actual = None
        if level == 1:
            m = RE_H1.match(text)
            if m:
                actual = zh_to_num(m.group(0).replace("、", ""))
        elif level == 2:
            m = RE_H2.match(text)
            if m:
                actual = zh_to_num(m.group(0).replace("（", "").replace("）", ""))
        elif level == 3:
            m = RE_H3.match(text)
            if m:
                actual = int(m.group(0).replace(".", ""))
        else:
            m = RE_H4.match(text)
            if m:
                actual = int(m.group(0).replace("（", "").replace("）", ""))

        if actual is not None and actual != expected[level]:
            warnings.append(
                f"第{i+1}个标题编号疑似跳号/混用：层级 H{level} 当前 {actual}，期望 {expected[level]}"
            )

    return warnings


def import_docx(file_bytes: bytes, preserve_formatting: bool = True) -> tuple[dict[str, Any], dict[str, Any], dict[str, Any]]:
    doc = Document(io.BytesIO(file_bytes))

    nodes: list[dict[str, Any]] = []
    unrecognized_title_count = 0
    headings: list[tuple[int, str]] = []
    table_warnings: list[str] = []
    extracted_title = ""
    title_consumed = False
    extracted_title_attrs: dict[str, Any] | None = None

    for paragraph_index, p in enumerate(doc.paragraphs):
        text = (p.text or "").strip()
        if not text:
            continue

        if not extracted_title and paragraph_index <= 4 and _looks_like_main_title(text, p):
            extracted_title = text
            title_consumed = True
            if preserve_formatting:
                extracted_title_attrs = _extract_paragraph_attrs(p)
            continue

        level = _detect_level(text, p)
        if level:
            headings.append((level, text))
            node_attrs = _extract_paragraph_attrs(p) if preserve_formatting else {}
            nodes.append(_make_heading(level, text, node_attrs))
        else:
            if preserve_formatting:
                node_attrs = _extract_paragraph_attrs(p)
            else:
                node_attrs = {"firstLineIndentChars": 2}
            nodes.append(_make_paragraph(text, node_attrs))
            if _looks_like_heading_candidate(text):
                unrecognized_title_count += 1

    for idx, table in enumerate(doc.tables):
        try:
            nodes.append(_table_to_node(table))
        except Exception as exc:  # pragma: no cover
            table_warnings.append(f"表格 {idx+1} 解析失败: {exc}")

    body = {"type": "doc", "content": nodes}

    structured = {
        "title": extracted_title,
        "mainTo": "",
        "signOff": "",
        "docNo": "",
        "signatory": "",
        "copyNo": "",
        "date": "",
        "exportWithRedhead": False,
        "attachments": [],
        "importedTitleAttrs": extracted_title_attrs,
    }

    # 从前几段尝试抽取文号并做括号归一
    for node in nodes[:8]:
        if node.get("type") in {"heading", "paragraph"}:
            t = node.get("content", [{}])[0].get("text", "")
            if re.search(r"\d{4}", t) and ("号" in t or "文" in t):
                structured["docNo"] = normalize_doc_no_brackets(t)
                break

    report = {
        "unrecognizedTitleCount": unrecognized_title_count,
        "numberingWarnings": _parse_numbering_warning(headings),
        "tableWarnings": table_warnings,
        "notes": [
            "导入时已忽略原 DOCX 页眉/红头（按系统红头模板重建）。",
            "已尽量保留原文段落的字体、字号、行距与缩进。",
        ],
    }
    if title_consumed:
        report["notes"].insert(0, "已识别并提取文档主标题。")

    return body, structured, report
