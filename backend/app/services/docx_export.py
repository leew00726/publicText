from __future__ import annotations

import io
import os
import re
from collections import defaultdict
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.services.checker import normalize_doc_no_brackets

RE_SUFFIX_LINE = re.compile(
    r"^(主\s*持(?:\s*人|\s*者)?|参\s*(?:加|会)(?:\s*人|\s*人员|\s*名单)?|列\s*席(?:\s*人|\s*人员)?|出\s*席(?:\s*人|\s*人员)?|记\s*录(?:\s*人|\s*员)?|发\s*(?:送|至|文)|主\s*送|抄\s*送|分\s*送)\s*[：:]"
)
RE_SUFFIX_LINE_CAPTURE = re.compile(
    r"^((?:主\s*持(?:\s*人|\s*者)?|参\s*(?:加|会)(?:\s*人|\s*人员|\s*名单)?|列\s*席(?:\s*人|\s*人员)?|出\s*席(?:\s*人|\s*人员)?|记\s*录(?:\s*人|\s*员)?|发\s*(?:送|至|文)|主\s*送|抄\s*送|分\s*送)\s*[：:])(\s*.*)$"
)
RE_DOC_ISSUE_LINE = re.compile(r"^(?:\d{4}\s*年第\s*[0-9一二三四五六七八九十百千]+\s*期|第\s*[0-9一二三四五六七八九十百千]+\s*期)$")
RE_HEADER_SIGNATORY = re.compile(r"签发人\s*[：:]")
RE_TITLE_KEYWORD = re.compile(
    r"(报告|纪要|请示|函|通知|方案|总结|通报|决定|公告|意见|办法|细则|规定|计划|说明|简报|要点|清单|材料)$"
)
RE_REFERENCE_DOCNO_FIRST = re.compile(
    r"(?P<docno>[A-Za-z0-9\u4e00-\u9fa5〔〕\-\u2014]+〔\d{2,4}〕[0-9一二三四五六七八九十百千]+号)\s*[《〈](?P<title>[^》〉]+)[》〉]"
)
REFERENCE_LEAD_PREFIXES = (
    "请参照",
    "参照",
    "按照",
    "根据",
    "依照",
    "依据",
    "遵照",
    "对照",
    "参见",
)


def _set_run_font(run, family: str, size_pt: float, bold: bool = False, color_hex: str | None = None):
    run.font.name = family
    run._element.rPr.rFonts.set(qn("w:eastAsia"), family)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_hex:
        hex_color = color_hex.replace("#", "")
        if len(hex_color) == 6:
            run.font.color.rgb = RGBColor.from_string(hex_color)


def _insert_page_number(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def _add_paragraph_border(paragraph, color: str = "D40000", size_eighth_point: int = 12):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size_eighth_point))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.replace("#", ""))
    pbdr.append(bottom)


def _node_text(node: dict[str, Any]) -> str:
    if not node:
        return ""
    if node.get("type") == "text":
        return node.get("text", "")
    return "".join(_node_text(c) for c in (node.get("content") or []))


def _strip_attachment_ext(name: str) -> str:
    return os.path.splitext(name)[0]


def _format_zh_date(date_text: str) -> str:
    value = (date_text or "").strip()
    if not value:
        return ""
    m = re.match(r"^(\d{4})-(\d{2})-(\d{2})$", value)
    if not m:
        return value
    year = int(m.group(1))
    month = int(m.group(2))
    day = int(m.group(3))
    return f"{year}年{month}月{day}日"


def _normalize_line_count(value: Any, default: int) -> int:
    try:
        if value is None:
            raise ValueError
        return max(int(round(float(value))), 0)
    except (TypeError, ValueError):
        return default


def _format_attachment_name(name: str, use_book_title_marks: bool = False) -> str:
    value = os.path.splitext(str(name or ""))[0].strip()
    value = value.strip("《》〈〉")
    value = re.sub(r"[。；，、.!?！？：:;]+$", "", value)
    if not value:
        return ""
    if use_book_title_marks:
        return f"《{value}》"
    return value


def _format_attachment_prefix(index: Any, item_suffix_punctuation: str = "dot") -> str:
    idx = str(index or "").strip() or "1"
    return f"{idx}. " if item_suffix_punctuation == "dot" else f"{idx} "


def _normalize_reference_text(text: str, reference_rules: dict[str, Any]) -> str:
    value = str(text or "")
    if not value:
        return value

    if str(reference_rules.get("yearBrackets") or "").strip() == "〔〕":
        value = normalize_doc_no_brackets(value)

    if str(reference_rules.get("citationOrder") or "").strip() == "titleThenDocNo":
        def _replace(match: re.Match[str]) -> str:
            docno = match.group("docno")
            lead_prefix = ""
            for prefix in REFERENCE_LEAD_PREFIXES:
                if docno.startswith(prefix) and len(docno) > len(prefix):
                    lead_prefix = prefix
                    docno = docno[len(prefix) :]
                    break
            return f"{lead_prefix}《{match.group('title')}》（{docno}）"

        value = RE_REFERENCE_DOCNO_FIRST.sub(_replace, value)
    return value


def _resolve_title_lines(title: str, arrangement: str | None) -> list[str]:
    raw = str(title or "").strip()
    if not raw:
        return []
    compact = re.sub(r"\s+", "", raw)
    if arrangement != "trapezoid" or len(compact) < 8:
        return [raw]

    best_lines: list[str] = [compact]
    best_score = float("inf")
    total = len(compact)
    candidate_line_counts = [2] if total < 12 else [2, 3]

    for line_count in candidate_line_counts:
        if total < line_count * 3:
            continue

        if line_count == 2:
            for first_len in range(3, total - 2):
                lengths = [first_len, total - first_len]
                if lengths[0] > lengths[1]:
                    continue
                score = (lengths[1] - lengths[0]) * 4 + abs(lengths[0] - total / 2) * 2
                if score < best_score:
                    best_score = score
                    best_lines = [compact[:first_len], compact[first_len:]]
        else:
            for first_len in range(3, total - 5):
                for second_len in range(first_len, total - first_len - 2):
                    third_len = total - first_len - second_len
                    lengths = [first_len, second_len, third_len]
                    if third_len < second_len:
                        continue
                    score = (lengths[-1] - lengths[0]) * 4 + sum(abs(length - total / 3) for length in lengths)
                    if score < best_score:
                        best_score = score
                        best_lines = [
                            compact[:first_len],
                            compact[first_len : first_len + second_len],
                            compact[first_len + second_len :],
                        ]

    return best_lines


def _apply_attachment_paragraph_style(
    paragraph,
    indent_chars: float,
    font_size_pt: float,
    prefix_text: str,
    wrap_align: str,
):
    paragraph.paragraph_format.line_spacing = Pt(28)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    if wrap_align == "text":
        prefix_width_pt = max(len(prefix_text), 1) * font_size_pt
        paragraph.paragraph_format.left_indent = Pt((indent_chars * font_size_pt) + prefix_width_pt)
        paragraph.paragraph_format.first_line_indent = Pt(-prefix_width_pt)
        return

    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(indent_chars * font_size_pt)


def _extract_bind_map(structured_fields: dict[str, Any], unit_name: str) -> dict[str, str]:
    return {
        "unitName": unit_name,
        "docNo": normalize_doc_no_brackets(structured_fields.get("docNo", "")),
        "signatory": structured_fields.get("signatory", ""),
        "copyNo": structured_fields.get("copyNo", ""),
        "fixedText": "",
    }


def _group_elements_by_y(elements: list[dict[str, Any]], tolerance: float = 0.05) -> list[list[dict[str, Any]]]:
    enabled = [e for e in elements if e.get("enabled", True)]
    enabled.sort(key=lambda x: float(x.get("yCm", 0)))

    groups: list[list[dict[str, Any]]] = []
    for elem in enabled:
        y = float(elem.get("yCm", 0))
        if not groups:
            groups.append([elem])
            continue
        prev_y = float(groups[-1][0].get("yCm", 0))
        if abs(y - prev_y) <= tolerance:
            groups[-1].append(elem)
        else:
            groups.append([elem])
    return groups


def _apply_header_template(doc: Document, template: dict[str, Any], structured_fields: dict[str, Any], unit_name: str):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.header_distance = Cm(0)

    header = section.first_page_header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    bind_map = _extract_bind_map(structured_fields, unit_name)
    groups = _group_elements_by_y(template.get("elements", []))
    content_width_cm = 21 - 2.7 - 2.5

    prev_y = 0.0
    prev_height = 0.0

    for group in groups:
        y_cm = float(group[0].get("yCm", 0))
        paragraph = header.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        delta = max(y_cm - (prev_y + prev_height), 0)
        paragraph.paragraph_format.space_before = Cm(delta)

        has_line = any(e.get("type") == "line" for e in group)
        text_elems = [e for e in group if e.get("type") == "text"]

        doc_no = next((e for e in text_elems if e.get("bind") == "docNo"), None)
        signatory = next((e for e in text_elems if e.get("bind") == "signatory"), None)

        if has_line and not text_elems:
            _add_paragraph_border(paragraph, color="D40000", size_eighth_point=12)
            paragraph.paragraph_format.space_before = Cm(delta)
            prev_height = 0.08
            prev_y = y_cm
            continue

        if doc_no and signatory:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            tab_stops = paragraph.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Cm(content_width_cm), WD_TAB_ALIGNMENT.RIGHT)
            left_text = bind_map.get("docNo", "")
            right_text = bind_map.get("signatory", "")
            r1 = paragraph.add_run(left_text)
            _set_run_font(r1, "仿宋_GB2312", 16)
            paragraph.add_run("\t")
            r2 = paragraph.add_run(right_text)
            _set_run_font(r2, "仿宋_GB2312", 16)
            prev_height = (16 / 72) * 2.54 * 1.2
            prev_y = y_cm
            continue

        for elem in text_elems:
            bind = elem.get("bind")
            value = elem.get("fixedText") if bind == "fixedText" else bind_map.get(bind, "")
            if not value and not elem.get("visibleIfEmpty", False):
                continue

            align = ((elem.get("text") or {}).get("align")) or "left"
            anchor = ((elem.get("x") or {}).get("anchor")) or "marginLeft"
            if align == "center" or anchor == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == "right" or anchor == "marginRight":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            font_conf = ((elem.get("text") or {}).get("font")) or {}
            family = font_conf.get("family", "仿宋_GB2312")
            size_pt = float(font_conf.get("sizePt", 16))
            bold = bool(font_conf.get("bold", False))
            color = font_conf.get("color", "#000000")

            run = paragraph.add_run(str(value))
            _set_run_font(run, family, size_pt, bold=bold, color_hex=color)
            prev_height = (size_pt / 72) * 2.54 * 1.2

        if has_line:
            _add_paragraph_border(paragraph, color="D40000", size_eighth_point=12)

        prev_y = y_cm


def _safe_float(value: Any, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _safe_float_or_none(value: Any) -> float | None:
    try:
        if value is None:
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _normalize_color_hex(value: Any) -> str | None:
    if value is None:
        return None
    raw = str(value).strip()
    if not raw:
        return None
    if raw.startswith("#"):
        raw = raw[1:]
    if len(raw) != 6:
        return None
    return f"#{raw.upper()}"


def _resolve_topic_style_rules(structured_fields: dict[str, Any]) -> tuple[dict[str, Any], dict[str, Any]]:
    rules = structured_fields.get("topicTemplateRules")
    if not isinstance(rules, dict):
        return {}, {}

    body = rules.get("body") if isinstance(rules.get("body"), dict) else {}
    headings = rules.get("headings") if isinstance(rules.get("headings"), dict) else {}
    return body, headings


def _looks_like_fixed_title_node(node: dict[str, Any]) -> bool:
    text = _node_text(node).replace("\u00A0", " ").strip()
    if not text:
        return False
    if RE_HEADER_SIGNATORY.search(text) or RE_DOC_ISSUE_LINE.match(text) or _is_suffix_line(text):
        return False
    if any(keyword in text for keyword in ["有限公司", "有限责任公司", "集团", "公司", "委员会", "办公室", "政府"]):
        return False
    return bool(RE_TITLE_KEYWORD.search(text))


def _has_fixed_title_content(structured_fields: dict[str, Any]) -> bool:
    rules = structured_fields.get("topicTemplateRules")
    if not isinstance(rules, dict):
        return False
    content_template = rules.get("contentTemplate")
    if not isinstance(content_template, dict):
        return False
    if content_template.get("titleMode") == "fixed":
        return True
    if content_template.get("titleMode") == "dynamic":
        return False
    leading_nodes = content_template.get("leadingNodes")
    if not isinstance(leading_nodes, list):
        return False
    return any(isinstance(node, dict) and _looks_like_fixed_title_node(node) for node in leading_nodes)


def _apply_body_paragraph_style(paragraph, body_style: dict[str, Any]):
    paragraph.paragraph_format.line_spacing = Pt(_safe_float(body_style.get("lineSpacingPt"), 28))
    first_indent_pt = _safe_float_or_none(body_style.get("firstLineIndentPt"))
    if first_indent_pt is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first_indent_pt)
    else:
        first_indent_chars = _safe_float_or_none(body_style.get("firstLineIndentChars"))
        paragraph.paragraph_format.first_line_indent = Pt((first_indent_chars if first_indent_chars is not None else 2) * 16)
    paragraph.paragraph_format.space_before = Pt(_safe_float(body_style.get("spaceBeforePt"), 0))
    paragraph.paragraph_format.space_after = Pt(_safe_float(body_style.get("spaceAfterPt"), 0))


def _apply_heading_style(paragraph, level: int, heading_styles: dict[str, Any]):
    level_key = f"level{level}"
    level_style = heading_styles.get(level_key) if isinstance(heading_styles.get(level_key), dict) else {}

    paragraph.paragraph_format.line_spacing = Pt(_safe_float(level_style.get("lineSpacingPt"), 28))
    first_indent_pt = _safe_float_or_none(level_style.get("firstLineIndentPt"))
    if first_indent_pt is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first_indent_pt)
    else:
        first_indent_chars = _safe_float_or_none(level_style.get("firstLineIndentChars"))
        paragraph.paragraph_format.first_line_indent = Pt((first_indent_chars if first_indent_chars is not None else 2) * 16)
    paragraph.paragraph_format.space_before = Pt(_safe_float(level_style.get("spaceBeforePt"), 0))
    paragraph.paragraph_format.space_after = Pt(_safe_float(level_style.get("spaceAfterPt"), 0))

    default_style_map = {
        1: ("黑体", 16, False),
        2: ("黑体", 16, False),
        3: ("仿宋_GB2312", 16, False),
        4: ("仿宋_GB2312", 16, False),
    }
    default_family, default_size, default_bold = default_style_map.get(level, ("仿宋_GB2312", 16, False))

    family = level_style.get("fontFamily") or default_family
    size = _safe_float(level_style.get("fontSizePt"), default_size)
    bold = bool(level_style.get("bold", default_bold))
    return family, size, bold


def _resolve_node_attrs(node: dict[str, Any]) -> dict[str, Any]:
    attrs = node.get("attrs")
    if isinstance(attrs, dict):
        return attrs
    return {}


def _apply_node_paragraph_overrides(paragraph, node_attrs: dict[str, Any], default_font_size_pt: float):
    align = str(node_attrs.get("textAlign") or "").strip().lower()
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    line_spacing_pt = _safe_float_or_none(node_attrs.get("lineSpacingPt"))
    if line_spacing_pt is not None:
        paragraph.paragraph_format.line_spacing = Pt(line_spacing_pt)

    space_before_pt = _safe_float_or_none(node_attrs.get("spaceBeforePt"))
    if space_before_pt is not None:
        paragraph.paragraph_format.space_before = Pt(space_before_pt)

    space_after_pt = _safe_float_or_none(node_attrs.get("spaceAfterPt"))
    if space_after_pt is not None:
        paragraph.paragraph_format.space_after = Pt(space_after_pt)

    left_indent_pt = _safe_float_or_none(node_attrs.get("leftIndentPt"))
    if left_indent_pt is not None:
        paragraph.paragraph_format.left_indent = Pt(left_indent_pt)

    right_indent_pt = _safe_float_or_none(node_attrs.get("rightIndentPt"))
    if right_indent_pt is not None:
        paragraph.paragraph_format.right_indent = Pt(right_indent_pt)

    first_indent_pt = _safe_float_or_none(node_attrs.get("firstLineIndentPt"))
    if first_indent_pt is not None:
        paragraph.paragraph_format.first_line_indent = Pt(first_indent_pt)
    else:
        first_indent_chars = _safe_float_or_none(node_attrs.get("firstLineIndentChars"))
        if first_indent_chars is not None:
            paragraph.paragraph_format.first_line_indent = Pt(first_indent_chars * default_font_size_pt)
        elif align in {"center", "right"}:
            paragraph.paragraph_format.first_line_indent = Pt(0)


def _resolve_node_run_style(node_attrs: dict[str, Any], fallback_family: str, fallback_size_pt: float, fallback_bold: bool):
    family = str(node_attrs.get("fontFamily") or fallback_family)
    size = _safe_float(node_attrs.get("fontSizePt"), fallback_size_pt)
    node_bold = node_attrs.get("bold")
    bold = bool(node_bold) if isinstance(node_bold, bool) else fallback_bold
    color_hex = _normalize_color_hex(node_attrs.get("colorHex"))
    return family, size, bold, color_hex


def _append_red_divider_paragraph(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(1)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    _add_paragraph_border(paragraph, color="D40000", size_eighth_point=12)


def _is_suffix_line(text: str) -> bool:
    return bool(RE_SUFFIX_LINE.match((text or "").strip()))


def _normalize_suffix_line_attrs(node_attrs: dict[str, Any], body_style: dict[str, Any]) -> dict[str, Any]:
    normalized = dict(node_attrs or {})
    body_family = body_style.get("fontFamily")
    if isinstance(body_family, str) and body_family.strip():
        normalized["fontFamily"] = body_family.strip()

    body_size = _safe_float_or_none(body_style.get("fontSizePt"))
    if body_size is not None:
        normalized["fontSizePt"] = body_size

    body_line = _safe_float_or_none(body_style.get("lineSpacingPt"))
    if body_line is not None:
        normalized["lineSpacingPt"] = body_line

    body_indent_pt = _safe_float_or_none(body_style.get("firstLineIndentPt"))
    body_indent_chars = _safe_float_or_none(body_style.get("firstLineIndentChars"))
    if body_indent_pt is not None:
        normalized["firstLineIndentPt"] = body_indent_pt
        normalized.pop("firstLineIndentChars", None)
    elif body_indent_chars is not None:
        normalized["firstLineIndentChars"] = body_indent_chars
        normalized.pop("firstLineIndentPt", None)

    normalized["textAlign"] = "left"
    normalized["bold"] = False
    return normalized


def _split_suffix_line(text: str) -> tuple[str | None, str]:
    value = (text or "").strip()
    if not value:
        return None, ""
    match = RE_SUFFIX_LINE_CAPTURE.match(value)
    if not match:
        return None, value
    label = match.group(1)
    rest = match.group(2) or ""
    return label, rest


def _iter_nodes(doc_json: dict[str, Any]) -> list[dict[str, Any]]:
    return list((doc_json or {}).get("content") or [])


def export_docx(
    document_data: dict[str, Any],
    unit_name: str,
    redhead_template: dict[str, Any],
    include_redhead: bool = True,
) -> bytes:
    doc = Document()
    section = doc.sections[0]

    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.5)

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _insert_page_number(footer_para)

    structured_fields = document_data.get("structuredFields", {})
    body_style, heading_styles = _resolve_topic_style_rules(structured_fields)
    topic_template_rules = structured_fields.get("topicTemplateRules") if isinstance(structured_fields, dict) else {}
    if not isinstance(topic_template_rules, dict):
        topic_template_rules = {}
    title_style = topic_template_rules.get("title")
    if not isinstance(title_style, dict):
        title_style = {}
    imported_title_style = structured_fields.get("importedTitleAttrs")
    if not isinstance(imported_title_style, dict):
        imported_title_style = {}
    references_rules = topic_template_rules.get("references")
    if not isinstance(references_rules, dict):
        references_rules = {}
    attachments_rules = topic_template_rules.get("attachments")
    if not isinstance(attachments_rules, dict):
        attachments_rules = {}
    signature_rules = topic_template_rules.get("signature")
    if not isinstance(signature_rules, dict):
        signature_rules = {}
    suppress_auto_frontmatter = _has_fixed_title_content(structured_fields)
    if include_redhead:
        _apply_header_template(doc, redhead_template, structured_fields, unit_name)

    title = ""
    if not suppress_auto_frontmatter:
        title = structured_fields.get("title") or document_data.get("title") or ""
    if title:
        title_lines = _resolve_title_lines(title, str(title_style.get("arrangement") or "").strip())
        title_align = str(title_style.get("textAlign") or imported_title_style.get("textAlign") or "center").strip().lower()
        title_alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(title_align, WD_ALIGN_PARAGRAPH.CENTER)

        for title_line in title_lines:
            p_title = doc.add_paragraph()
            p_title.alignment = title_alignment
            p_title.paragraph_format.line_spacing = Pt(
                _safe_float(title_style.get("lineSpacingPt"), _safe_float(imported_title_style.get("lineSpacingPt"), 28))
            )
            p_title.paragraph_format.space_before = Pt(0)
            p_title.paragraph_format.space_after = Pt(0)
            run = p_title.add_run(title_line)
            _set_run_font(
                run,
                str(title_style.get("fontFamily") or imported_title_style.get("fontFamily") or "方正小标宋简"),
                _safe_float(title_style.get("fontSizePt"), _safe_float(imported_title_style.get("fontSizePt"), 22)),
                bool(title_style.get("bold", imported_title_style.get("bold", False))),
                color_hex=_normalize_color_hex(title_style.get("colorHex") or imported_title_style.get("colorHex")),
            )

    main_to = ""
    if not suppress_auto_frontmatter:
        main_to = structured_fields.get("mainTo", "").strip()
    if main_to:
        p_main = doc.add_paragraph()
        p_main.paragraph_format.line_spacing = Pt(28)
        p_main.paragraph_format.first_line_indent = Pt(0)
        run = p_main.add_run(main_to)
        _set_run_font(run, "仿宋_GB2312", 16)

    in_suffix_block = False
    for node in _iter_nodes(document_data.get("body", {})):
        ntype = node.get("type")
        node_attrs = _resolve_node_attrs(node)

        if ntype == "heading":
            level = int((node.get("attrs") or {}).get("level", 1))
            p = doc.add_paragraph()
            family, size, bold = _apply_heading_style(p, level, heading_styles)
            family, size, bold, color_hex = _resolve_node_run_style(node_attrs, family, size, bold)
            _apply_node_paragraph_overrides(p, node_attrs, size)
            text = _normalize_reference_text(_node_text(node), references_rules)
            r = p.add_run(text)
            _set_run_font(r, family, size, bold, color_hex=color_hex)
            continue

        if ntype == "paragraph":
            if bool(node_attrs.get("dividerRed")):
                _append_red_divider_paragraph(doc)
                continue

            p = doc.add_paragraph()
            _apply_body_paragraph_style(p, body_style)
            default_family = body_style.get("fontFamily", "仿宋_GB2312")
            default_size = _safe_float(body_style.get("fontSizePt"), 16)
            default_bold = bool(body_style.get("bold", False))
            text = _normalize_reference_text(_node_text(node), references_rules)
            if _is_suffix_line(text):
                in_suffix_block = True
            if in_suffix_block and text.strip():
                node_attrs = _normalize_suffix_line_attrs(node_attrs, body_style)
            _apply_node_paragraph_overrides(p, node_attrs, default_size)
            family, size, bold, color_hex = _resolve_node_run_style(node_attrs, default_family, default_size, default_bold)
            label_text, body_text = _split_suffix_line(text)
            if label_text:
                label_run = p.add_run(label_text)
                _set_run_font(label_run, "黑体", size, False)
                if body_text:
                    body_run = p.add_run(body_text)
                    _set_run_font(body_run, family, size, bold, color_hex=color_hex)
            else:
                r = p.add_run(text)
                _set_run_font(r, family, size, bold, color_hex=color_hex)
            continue

        if ntype == "table":
            rows = node.get("content") or []
            matrix: list[list[str]] = []
            for row in rows:
                row_cells = []
                for cell in (row.get("content") or []):
                    row_cells.append(_node_text(cell).strip())
                if row_cells:
                    matrix.append(row_cells)

            if not matrix:
                continue

            col_count = max(len(r) for r in matrix)
            table = doc.add_table(rows=len(matrix), cols=col_count)
            table.autofit = False
            for r_idx, row in enumerate(matrix):
                for c_idx in range(col_count):
                    value = row[c_idx] if c_idx < len(row) else ""
                    cell_para = table.cell(r_idx, c_idx).paragraphs[0]
                    cell_para.paragraph_format.line_spacing = Pt(28)
                    rr = cell_para.add_run(value)
                    _set_run_font(rr, "仿宋_GB2312", 16)

    sign_off = (structured_fields.get("signOff", "") or "").strip()
    date_text = _format_zh_date(structured_fields.get("date", ""))
    if sign_off or date_text:
        sign_spacing_before_lines = _normalize_line_count(signature_rules.get("spacingBeforeLines"), 2)
        for _ in range(sign_spacing_before_lines):
            p_blank = doc.add_paragraph()
            p_blank.paragraph_format.line_spacing = Pt(28)
            p_blank.paragraph_format.first_line_indent = Pt(0)
            p_blank.paragraph_format.space_before = Pt(0)
            p_blank.paragraph_format.space_after = Pt(0)

        if sign_off:
            p_sign = doc.add_paragraph()
            p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_sign.paragraph_format.line_spacing = Pt(28)
            p_sign.paragraph_format.first_line_indent = Pt(0)
            r_sign = p_sign.add_run(sign_off)
            _set_run_font(r_sign, "仿宋_GB2312", 16)

        if date_text:
            p_date = doc.add_paragraph()
            p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_date.paragraph_format.line_spacing = Pt(28)
            p_date.paragraph_format.first_line_indent = Pt(0)
            r_date = p_date.add_run(date_text)
            _set_run_font(r_date, "仿宋_GB2312", 16)

    attachments = structured_fields.get("attachments") or []
    if attachments:
        attachments_spacing_before_lines = _normalize_line_count(attachments_rules.get("spacingBeforeLines"), 1)
        attachment_indent_chars = _safe_float(attachments_rules.get("indentChars"), 2)
        attachment_item_suffix = str(attachments_rules.get("itemSuffixPunctuation") or "dot").strip() or "dot"
        attachment_wrap_align = str(attachments_rules.get("wrapAlign") or "indent").strip() or "indent"
        attachment_use_book_title_marks = attachments_rules.get("useBookTitleMarks") is True
        for _ in range(attachments_spacing_before_lines):
            doc.add_paragraph("")
        attach_label = doc.add_paragraph()
        attach_label.paragraph_format.line_spacing = Pt(28)
        attach_label.paragraph_format.first_line_indent = Pt(attachment_indent_chars * 16)
        r = attach_label.add_run("附件：")
        _set_run_font(r, "仿宋_GB2312", 16)

        for item in attachments:
            idx = item.get("index")
            name = _format_attachment_name(item.get("name", ""), attachment_use_book_title_marks)
            prefix_text = _format_attachment_prefix(idx, attachment_item_suffix)
            p_attach = doc.add_paragraph()
            _apply_attachment_paragraph_style(
                p_attach,
                attachment_indent_chars,
                16,
                prefix_text,
                attachment_wrap_align,
            )
            rr = p_attach.add_run(f"{prefix_text}{name}")
            _set_run_font(rr, "仿宋_GB2312", 16)

        for item in attachments:
            doc.add_page_break()
            idx = item.get("index")
            title_text = _format_attachment_name(item.get("name", ""), attachment_use_book_title_marks)

            p_mark = doc.add_paragraph()
            p_mark.paragraph_format.line_spacing = Pt(28)
            p_mark.paragraph_format.first_line_indent = Pt(0)
            r_mark = p_mark.add_run(f"附件{idx}")
            _set_run_font(r_mark, "黑体", 16)

            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.line_spacing = Pt(28)
            r_title = p_title.add_run(title_text)
            _set_run_font(r_title, "方正小标宋简", 22)

            p_body = doc.add_paragraph()
            _apply_body_paragraph_style(p_body, body_style)
            r_body = p_body.add_run("（附件正文请在此处编辑）")
            _set_run_font(r_body, "仿宋_GB2312", 16)

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
