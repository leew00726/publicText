import io
import unittest

from docx import Document
from docx.oxml.ns import qn

from app.services.docx_export import export_docx


def _paragraph_texts(raw: bytes) -> list[str]:
    doc = Document(io.BytesIO(raw))
    return [p.text for p in doc.paragraphs if p.text]


class DocxExportTests(unittest.TestCase):
    def test_export_suppresses_auto_title_when_topic_has_fixed_leading_nodes(self) -> None:
        payload = {
            "title": "自动标题-不应出现",
            "structuredFields": {
                "title": "",
                "mainTo": "",
                "topicTemplateRules": {
                    "contentTemplate": {
                        "leadingNodes": [{"type": "paragraph", "content": [{"type": "text", "text": "固定头部"}]}],
                        "trailingNodes": [],
                        "bodyPlaceholder": "（请在此输入正文）",
                    }
                },
            },
            "body": {
                "type": "doc",
                "content": [{"type": "paragraph", "content": [{"type": "text", "text": "正文段落"}]}],
            },
        }

        raw = export_docx(payload, unit_name="测试单位", redhead_template={"elements": [], "page": {}}, include_redhead=False)
        texts = _paragraph_texts(raw)

        self.assertIn("正文段落", texts)
        self.assertNotIn("自动标题-不应出现", texts)

    def test_export_keeps_auto_title_without_fixed_leading_nodes(self) -> None:
        payload = {
            "title": "自动标题-应出现",
            "structuredFields": {"title": "", "mainTo": "", "topicTemplateRules": {}},
            "body": {"type": "doc", "content": []},
        }

        raw = export_docx(payload, unit_name="测试单位", redhead_template={"elements": [], "page": {}}, include_redhead=False)
        texts = _paragraph_texts(raw)

        self.assertIn("自动标题-应出现", texts)

    def test_export_normalizes_suffix_lines_to_body_font(self) -> None:
        payload = {
            "title": "测试",
            "structuredFields": {
                "title": "",
                "mainTo": "",
                "topicTemplateRules": {"body": {"fontFamily": "仿宋_GB2312", "fontSizePt": 16, "lineSpacingPt": 28, "firstLineIndentPt": 32}},
            },
            "body": {
                "type": "doc",
                "content": [
                    {
                        "type": "paragraph",
                        "attrs": {"fontFamily": "黑体", "fontSizePt": 16, "bold": True, "firstLineIndentChars": 2},
                        "content": [{"type": "text", "text": "主 持：金刚善"}],
                    },
                    {
                        "type": "paragraph",
                        "attrs": {"fontFamily": "仿宋", "fontSizePt": 16, "bold": False, "firstLineIndentChars": 2},
                        "content": [{"type": "text", "text": "王振宇、刘冬冬、徐国涛"}],
                    },
                    {
                        "type": "paragraph",
                        "attrs": {"fontFamily": "黑体", "fontSizePt": 16, "bold": True, "firstLineIndentChars": 2},
                        "content": [{"type": "text", "text": "参会人员：张三、李四"}],
                    }
                ],
            },
        }

        raw = export_docx(payload, unit_name="测试单位", redhead_template={"elements": [], "page": {}}, include_redhead=False)
        doc = Document(io.BytesIO(raw))
        host_paragraph = next((p for p in doc.paragraphs if p.text.strip() == "主 持：金刚善"), None)
        self.assertIsNotNone(host_paragraph)
        self.assertGreaterEqual(len(host_paragraph.runs), 2)
        self.assertEqual(host_paragraph.runs[0].text, "主 持：")
        self.assertEqual(host_paragraph.runs[0].font.name, "黑体")
        self.assertFalse(bool(host_paragraph.runs[0].bold))
        self.assertEqual(host_paragraph.runs[1].font.name, "仿宋_GB2312")
        self.assertFalse(bool(host_paragraph.runs[1].bold))

        continuation = next((p for p in doc.paragraphs if p.text.strip() == "王振宇、刘冬冬、徐国涛"), None)
        self.assertIsNotNone(continuation)
        self.assertTrue(continuation.runs)
        self.assertEqual(continuation.runs[0].font.name, "仿宋_GB2312")

        attendee = next((p for p in doc.paragraphs if p.text.strip() == "参会人员：张三、李四"), None)
        self.assertIsNotNone(attendee)
        self.assertGreaterEqual(len(attendee.runs), 2)
        self.assertEqual(attendee.runs[0].text, "参会人员：")
        self.assertEqual(attendee.runs[0].font.name, "黑体")
        self.assertEqual(attendee.runs[1].font.name, "仿宋_GB2312")

    def test_export_renders_divider_red_paragraph_as_bottom_border(self) -> None:
        payload = {
            "title": "测试",
            "structuredFields": {
                "title": "",
                "mainTo": "",
                "topicTemplateRules": {"body": {"fontFamily": "仿宋_GB2312", "fontSizePt": 16, "lineSpacingPt": 28, "firstLineIndentPt": 32}},
            },
            "body": {
                "type": "doc",
                "content": [
                    {"type": "paragraph", "content": [{"type": "text", "text": "综合管理部  2026年1月12日  签发人："}]},
                    {"type": "paragraph", "attrs": {"dividerRed": True}, "content": []},
                    {"type": "paragraph", "content": [{"type": "text", "text": "正文首段。"}]},
                ],
            },
        }

        raw = export_docx(payload, unit_name="测试单位", redhead_template={"elements": [], "page": {}}, include_redhead=False)
        doc = Document(io.BytesIO(raw))
        divider_found = False
        for paragraph in doc.paragraphs:
            pPr = paragraph._p.pPr
            if pPr is None:
                continue
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is None:
                continue
            bottom = pBdr.find(qn("w:bottom"))
            if bottom is None:
                continue
            color = (bottom.get(qn("w:color")) or "").upper()
            if color == "D40000":
                divider_found = True
                break

        self.assertTrue(divider_found)


if __name__ == "__main__":
    unittest.main()
