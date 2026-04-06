import io
import os
import unittest
import uuid
from pathlib import Path
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient


ROOT = Path(__file__).resolve().parents[2]
DB_PATH = ROOT / "test_topics_api.db"

os.environ["DATABASE_URL"] = f"sqlite:///{DB_PATH.as_posix()}"
os.environ["STORAGE_MODE"] = "local"
os.environ["EXPORT_DIR"] = str((ROOT / "test-storage").as_posix())

from app.database import SessionLocal, engine  # noqa: E402
from app.main import app  # noqa: E402
from app.models import DeletionAuditEvent, DocumentFile, Topic, TopicTemplate, TopicTemplateDraft  # noqa: E402
from app.routers.topics import _build_patch_from_instruction, _extract_font_name  # noqa: E402


def _docx_bytes(text: str = "示例正文") -> bytes:
    doc = Document()
    doc.add_heading("一、总体要求", level=1)
    doc.add_paragraph(text)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _docx_bytes_for_level3_font_test(body_font: str = "仿宋_GB2312", level3_font: str = "黑体") -> bytes:
    doc = Document()
    heading = doc.add_heading("1. 三级标题示例", level=3)
    for run in heading.runs:
        run.font.name = level3_font

    para = doc.add_paragraph("这是用于模板修订测试的正文。")
    for run in para.runs:
        run.font.name = body_font

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _fake_pdf_bytes() -> bytes:
    return b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\ntrailer\n<<>>\n%%EOF"


class TopicApiTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        if DB_PATH.exists():
            DB_PATH.unlink()
        cls._client_cm = TestClient(app)
        cls.client = cls._client_cm.__enter__()

    @classmethod
    def tearDownClass(cls) -> None:
        cls._client_cm.__exit__(None, None, None)
        engine.dispose()
        if DB_PATH.exists():
            DB_PATH.unlink()

    def test_topic_flow_endpoints(self) -> None:
        company_resp = self.client.post("/api/units", json={"name": f"topic_api_company_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "周例会纪要", "description": "每周例会模板"},
        )
        self.assertEqual(created.status_code, 200)
        topic = created.json()
        self.assertEqual(topic["status"], "active")
        topic_id = topic["id"]

        topic_detail = self.client.get(f"/api/topics/{topic_id}")
        self.assertEqual(topic_detail.status_code, 200)
        self.assertEqual(topic_detail.json()["id"], topic_id)

        listed = self.client.get("/api/topics", params={"companyId": company_id})
        self.assertEqual(listed.status_code, 200)
        self.assertTrue(any(item["id"] == topic_id for item in listed.json()))

        with SessionLocal() as session:
            before = session.query(DocumentFile).count()

        analyze = self.client.post(
            f"/api/topics/{topic_id}/analyze",
            files=[
                (
                    "files",
                    (
                        "sample-a.docx",
                        _docx_bytes("第一份样本"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    (
                        "sample-b.docx",
                        _docx_bytes("第二份样本"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(analyze.status_code, 200)
        draft = analyze.json()["draft"]
        self.assertEqual(draft["status"], "draft")
        self.assertIn("body", draft["inferredRules"])

        with SessionLocal() as session:
            after = session.query(DocumentFile).count()
        self.assertEqual(before, after)

        revised = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={"instruction": "正文改成宋体", "patch": {"body": {"fontFamily": "宋体"}}},
        )
        self.assertEqual(revised.status_code, 200)
        revised_draft = revised.json()
        self.assertEqual(revised_draft["version"], 2)
        self.assertEqual(revised_draft["inferredRules"]["body"]["fontFamily"], "宋体")

        confirmed = self.client.post(f"/api/topics/{topic_id}/confirm-template")
        self.assertEqual(confirmed.status_code, 200)
        template = confirmed.json()["template"]
        self.assertTrue(template["effective"])
        first_template_id = template["id"]
        first_template_version = template["version"]

        revised_again = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={"instruction": "正文改成仿宋_GB2312", "patch": {"body": {"fontFamily": "仿宋_GB2312"}}},
        )
        self.assertEqual(revised_again.status_code, 200)

        confirmed_again = self.client.post(f"/api/topics/{topic_id}/confirm-template")
        self.assertEqual(confirmed_again.status_code, 200)

        templates = self.client.get(f"/api/topics/{topic_id}/templates")
        self.assertEqual(templates.status_code, 200)
        self.assertTrue(any(item["effective"] for item in templates.json()))

        audits = self.client.get(f"/api/topics/{topic_id}/audit-events")
        self.assertEqual(audits.status_code, 200)
        self.assertGreaterEqual(len(audits.json()), 1)
        first_audit = audits.json()[0]
        self.assertEqual(first_audit["topicId"], topic_id)
        self.assertIn("fileCount", first_audit)
        self.assertIn("totalBytes", first_audit)
        self.assertNotIn("filename", first_audit)
        self.assertNotIn("content", first_audit)

        effective = next(item for item in templates.json() if item["effective"])
        create_doc = self.client.post(
            f"/api/topics/{topic_id}/docs",
            json={"title": "周例会纪要（新建）", "topicTemplateId": first_template_id},
        )
        self.assertEqual(create_doc.status_code, 200)
        created_doc_id = create_doc.json()["id"]
        created_doc = self.client.get(f"/api/docs/{created_doc_id}")
        self.assertEqual(created_doc.status_code, 200)
        sf = created_doc.json()["structuredFields"]
        self.assertEqual(created_doc.json()["unitId"], company_id)
        self.assertEqual(created_doc.json()["docType"], "jiyao")
        self.assertEqual(sf["topicId"], topic_id)
        self.assertEqual(sf["topicTemplateId"], first_template_id)
        self.assertEqual(sf["topicTemplateVersion"], first_template_version)

    def test_topic_analyze_builds_dynamic_title_template_and_prefills_new_doc_title(self) -> None:
        company_resp = self.client.post("/api/units", json={"name": f"topic_title_tpl_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "报告题材", "description": "测试动态标题模板"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        def build_title_docx(title_text: str) -> bytes:
            doc = Document()
            company = doc.add_paragraph("华能云成数字产融科技（雄安）有限公司")
            company.alignment = 1
            company.runs[0].font.name = "方正小标宋简体"

            title = doc.add_paragraph(title_text)
            title.alignment = 1
            title.runs[0].font.name = "方正小标宋简体"

            issue = doc.add_paragraph("2026年第6期")
            issue.alignment = 1
            issue.runs[0].font.name = "黑体"

            body = doc.add_paragraph("这是正文第一段。")
            body.runs[0].font.name = "仿宋_GB2312"

            bio = io.BytesIO()
            doc.save(bio)
            return bio.getvalue()

        analyze = self.client.post(
            f"/api/topics/{topic_id}/analyze",
            files=[
                (
                    "files",
                    (
                        "report-a.docx",
                        build_title_docx("2026年度报告"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
                (
                    "files",
                    (
                        "report-b.docx",
                        build_title_docx("2025年度报告"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(analyze.status_code, 200)
        draft_rules = analyze.json()["draft"]["inferredRules"]
        self.assertEqual((draft_rules.get("title") or {}).get("templateText"), "{{变量1}}年度报告")

        confirm = self.client.post(f"/api/topics/{topic_id}/confirm-template")
        self.assertEqual(confirm.status_code, 200)

        create_doc = self.client.post(f"/api/topics/{topic_id}/docs", json={"title": "年度报告（新建）"})
        self.assertEqual(create_doc.status_code, 200)
        doc_id = create_doc.json()["id"]

        created_doc = self.client.get(f"/api/docs/{doc_id}")
        self.assertEqual(created_doc.status_code, 200)
        self.assertEqual(created_doc.json()["structuredFields"]["title"], "{{变量1}}年度报告")

    def test_revise_instruction_updates_level3_heading_font_without_overriding_body(self) -> None:
        company_resp = self.client.post("/api/units", json={"name": f"topic_heading_fix_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "标题修订测试题材", "description": "测试三级标题字体修订"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        analyze = self.client.post(
            f"/api/topics/{topic_id}/analyze",
            files=[
                (
                    "files",
                    (
                        "level3-font-sample.docx",
                        _docx_bytes_for_level3_font_test(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(analyze.status_code, 200)
        initial_rules = analyze.json()["draft"]["inferredRules"]
        initial_body_font = (initial_rules.get("body") or {}).get("fontFamily")

        revised = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={"instruction": "把三级标题改为宋体"},
        )
        self.assertEqual(revised.status_code, 200)
        revised_rules = revised.json()["inferredRules"]

        self.assertEqual(revised_rules["headings"]["level3"]["fontFamily"], "宋体")
        revised_body_font = (revised_rules.get("body") or {}).get("fontFamily")
        self.assertEqual(revised_body_font, initial_body_font)

    def test_revise_without_analyze_creates_initial_draft_from_instruction(self) -> None:
        company_resp = self.client.post("/api/units", json={"name": f"topic_text_only_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "纯文字训练题材", "description": "测试纯文字首版草稿"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        revised = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={"instruction": "正文字体统一改成宋体，主标题改为方正小标宋简体"},
        )
        self.assertEqual(revised.status_code, 200)
        revised_draft = revised.json()

        self.assertEqual(revised_draft["version"], 1)
        self.assertEqual(revised_draft["inferredRules"]["body"]["fontFamily"], "宋体")
        self.assertEqual(revised_draft["inferredRules"]["title"]["fontFamily"], "方正小标宋简")
        self.assertEqual(revised_draft["confidenceReport"], {})
        self.assertEqual(revised_draft["agentSummary"], "正文字体统一改成宋体，主标题改为方正小标宋简体")

    @patch("app.routers.topics.revise_topic_rules_with_deepseek")
    def test_revise_without_analyze_can_create_initial_draft_with_deepseek(self, mock_revise_with_deepseek) -> None:
        mock_revise_with_deepseek.return_value = {
            "patch": {"headings": {"level2": {"fontFamily": "楷体_GB2312"}}},
            "assistantReply": "已根据文字指令生成首版模板草稿。",
            "summary": "已生成首版模板草稿",
            "model": "deepseek-chat",
            "usage": {"total_tokens": 66},
        }

        company_resp = self.client.post("/api/units", json={"name": f"topic_text_ds_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "纯文字DeepSeek题材", "description": "测试纯文字 DeepSeek 首稿"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        revised = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={
                "instruction": "请直接根据要求生成首版模板草稿，二级标题改为楷体_GB2312",
                "useDeepSeek": True,
                "conversation": [{"role": "user", "content": "现在没有上传文件，请直接生成模板"}],
            },
        )
        self.assertEqual(revised.status_code, 200)
        revised_draft = revised.json()

        self.assertEqual(revised_draft["version"], 1)
        self.assertEqual(revised_draft["inferredRules"]["headings"]["level2"]["fontFamily"], "楷体_GB2312")
        self.assertEqual(revised_draft["agentSummary"], "已根据文字指令生成首版模板草稿。")
        self.assertTrue(mock_revise_with_deepseek.called)

    def test_build_patch_from_instruction_handles_title_and_numbering_fonts(self) -> None:
        instruction = (
            "一、 标题与正文字体规范 主标题：使用 方正小标宋简体 2 号。"
            "标题排列一般采用梯形或菱形，不宜只排一行。"
            "正文字体：统一使用 3 号仿宋_GB2312。"
            "二、 正文层级序号规则 第一层级：序号为“一、”，使用 黑体 3 号。"
            "第二层级：序号为“（一）”，使用 楷体_GB2312 3 号。"
            "第三层级：序号为“1.”，使用 仿宋_GB2312 3 号。"
            "第四层级：序号为“（1）”，使用 仿宋 3 号。"
        )

        self.assertIsNone(_extract_font_name("标题排列一般采用梯形或菱形，不宜只排一行。"))

        patch = _build_patch_from_instruction(instruction)

        self.assertEqual(patch["title"]["fontFamily"], "方正小标宋简")
        self.assertEqual(patch["body"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(patch["headings"]["level1"]["fontFamily"], "黑体")
        self.assertEqual(patch["headings"]["level2"]["fontFamily"], "楷体_GB2312")
        self.assertEqual(patch["headings"]["level3"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(patch["headings"]["level4"]["fontFamily"], "仿宋_GB2312")

    @patch("app.routers.topics.revise_topic_rules_with_deepseek")
    def test_revise_with_deepseek_ignores_invalid_font_family_from_model(self, mock_revise_with_deepseek) -> None:
        mock_revise_with_deepseek.return_value = {
            "patch": {"headings": {"level1": {"fontFamily": "梯形或菱形"}}},
            "assistantReply": "已根据指令更新模板规则。",
            "summary": "模板规则已更新",
            "model": "deepseek-chat",
            "usage": {"total_tokens": 21},
        }

        company_resp = self.client.post("/api/units", json={"name": f"topic_invalid_font_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "非法字体过滤题材", "description": "测试非法字体值过滤"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        analyze = self.client.post(
            f"/api/topics/{topic_id}/analyze",
            files=[
                (
                    "files",
                    (
                        "invalid-font-sample.docx",
                        _docx_bytes_for_level3_font_test(level3_font="黑体"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(analyze.status_code, 200)
        initial_level1_font = (
            ((analyze.json()["draft"]["inferredRules"] or {}).get("headings") or {}).get("level1") or {}
        ).get("fontFamily")

        revised = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={
                "instruction": "标题排列一般采用梯形或菱形，不宜只排一行。",
                "useDeepSeek": True,
            },
        )
        self.assertEqual(revised.status_code, 200)
        revised_level1_font = ((((revised.json()["inferredRules"] or {}).get("headings") or {}).get("level1") or {}).get(
            "fontFamily"
        ))

        self.assertEqual(revised_level1_font, initial_level1_font)
        self.assertNotEqual(revised_level1_font, "梯形或菱形")

    @patch("app.routers.topics.revise_topic_rules_with_deepseek")
    def test_revise_with_deepseek_chat_uses_model_patch(self, mock_revise_with_deepseek) -> None:
        mock_revise_with_deepseek.return_value = {
            "patch": {"headings": {"level3": {"fontFamily": "宋体"}}},
            "assistantReply": "已将三级标题字体调整为宋体。",
            "summary": "三级标题改为宋体",
            "model": "deepseek-chat",
            "usage": {"total_tokens": 88},
        }

        company_resp = self.client.post("/api/units", json={"name": f"topic_deepseek_fix_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "DeepSeek修订题材", "description": "测试DeepSeek修订"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        analyze = self.client.post(
            f"/api/topics/{topic_id}/analyze",
            files=[
                (
                    "files",
                    (
                        "deepseek-sample.docx",
                        _docx_bytes_for_level3_font_test(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(analyze.status_code, 200)

        revised = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={
                "instruction": "把三级标题改成宋体",
                "useDeepSeek": True,
                "conversation": [{"role": "user", "content": "请帮我统一三级标题字体"}],
            },
        )
        self.assertEqual(revised.status_code, 200)
        revised_draft = revised.json()
        self.assertEqual(revised_draft["inferredRules"]["headings"]["level3"]["fontFamily"], "宋体")
        self.assertEqual(revised_draft["agentSummary"], "已将三级标题字体调整为宋体。")
        self.assertTrue(mock_revise_with_deepseek.called)

    @patch("app.routers.topics.revise_topic_rules_with_deepseek")
    def test_revise_with_deepseek_chat_applies_instruction_fallback_patch(self, mock_revise_with_deepseek) -> None:
        mock_revise_with_deepseek.return_value = {
            "patch": {"body": {"lineSpacingPt": 28}},
            "assistantReply": "已根据指令调整模板规则。",
            "summary": "模板规则已更新",
            "model": "deepseek-chat",
            "usage": {"total_tokens": 88},
        }

        company_resp = self.client.post("/api/units", json={"name": f"topic_deepseek_fallback_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "DeepSeek兜底修订题材", "description": "测试DeepSeek兜底补丁"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        analyze = self.client.post(
            f"/api/topics/{topic_id}/analyze",
            files=[
                (
                    "files",
                    (
                        "deepseek-fallback-sample.docx",
                        _docx_bytes_for_level3_font_test(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(analyze.status_code, 200)

        revised = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={
                "instruction": "把三级标题改成宋体",
                "useDeepSeek": True,
                "conversation": [{"role": "user", "content": "请帮我统一三级标题字体"}],
            },
        )
        self.assertEqual(revised.status_code, 200)
        revised_rules = revised.json()["inferredRules"]
        self.assertEqual(revised_rules["headings"]["level3"]["fontFamily"], "宋体")
        self.assertEqual(revised_rules["body"]["lineSpacingPt"], 28)
        self.assertTrue(mock_revise_with_deepseek.called)

    @patch("app.routers.topics.revise_topic_rules_with_deepseek")
    def test_revise_with_deepseek_instruction_parser_fills_missing_advanced_rules(self, mock_revise_with_deepseek) -> None:
        mock_revise_with_deepseek.return_value = {
            "patch": {"title": {"fontFamily": "方正小标宋简"}},
            "assistantReply": "已根据指令更新模板规则。",
            "summary": "模板规则已更新",
            "model": "deepseek-chat",
            "usage": {"total_tokens": 96},
        }

        company_resp = self.client.post("/api/units", json={"name": f"topic_deepseek_parse_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "DeepSeek指令解析题材", "description": "测试高级规则指令解析"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        revised = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={
                "instruction": (
                    "标题使用方正小标宋简体2号，采用梯形排列并居中；"
                    "正文统一使用3号仿宋_GB2312字体，行间距设为固定值28磅，首行左空两格；"
                    "文中引用公文需先引标题后引发号，年份需使用六角括号〔 〕；"
                    "第一层标题为黑体3号（不加标点），第二层为楷体3号，第三、四层均为仿宋3号（必须加标点）；"
                    "附件部分需与正文空一行，左空两格起排，序号末尾不加标点，回行需与文字对齐，且附件名不加书名号；"
                    "最后，落款盖章处需与上文空两行。"
                ),
                "useDeepSeek": True,
            },
        )
        self.assertEqual(revised.status_code, 200)
        revised_rules = revised.json()["inferredRules"]

        self.assertEqual(revised_rules["title"]["fontFamily"], "方正小标宋简")
        self.assertEqual(revised_rules["title"]["fontSizePt"], 22)
        self.assertEqual(revised_rules["title"]["textAlign"], "center")
        self.assertEqual(revised_rules["title"]["arrangement"], "trapezoid")
        self.assertEqual(revised_rules["body"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(revised_rules["body"]["lineSpacingPt"], 28)
        self.assertEqual(revised_rules["body"]["firstLineIndentChars"], 2)
        self.assertEqual(revised_rules["headings"]["level1"]["fontFamily"], "黑体")
        self.assertEqual(revised_rules["headings"]["level1"]["punctuation"], False)
        self.assertEqual(revised_rules["headings"]["level2"]["fontFamily"], "楷体_GB2312")
        self.assertEqual(revised_rules["headings"]["level3"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(revised_rules["headings"]["level3"]["punctuation"], True)
        self.assertEqual(revised_rules["headings"]["level4"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(revised_rules["references"]["citationOrder"], "titleThenDocNo")
        self.assertEqual(revised_rules["references"]["yearBrackets"], "〔〕")
        self.assertEqual(revised_rules["attachments"]["spacingBeforeLines"], 1)
        self.assertEqual(revised_rules["attachments"]["indentChars"], 2)
        self.assertEqual(revised_rules["attachments"]["itemSuffixPunctuation"], "none")
        self.assertEqual(revised_rules["attachments"]["wrapAlign"], "text")
        self.assertFalse(revised_rules["attachments"]["useBookTitleMarks"])
        self.assertEqual(revised_rules["signature"]["spacingBeforeLines"], 2)
        self.assertTrue(mock_revise_with_deepseek.called)

    @patch("app.routers.topics.revise_topic_rules_with_deepseek")
    def test_revise_with_deepseek_normalizes_natural_language_patch_values(self, mock_revise_with_deepseek) -> None:
        mock_revise_with_deepseek.return_value = {
            "patch": {
                "title": {
                    "fontFamily": "方正小标宋简体",
                    "fontSize": "2号",
                    "textAlign": "居中",
                    "arrangement": "梯形",
                },
                "body": {
                    "fontFamily": "仿宋",
                    "fontSize": "3号",
                    "lineHeight": "固定值28磅",
                    "textIndent": "左空两格",
                },
                "headings": {
                    "level1": {"fontFamily": "黑体", "fontSize": "3号", "punctuation": False},
                    "level2": {"fontFamily": "楷体", "fontSize": "3号"},
                    "level3": {"fontFamily": "仿宋", "fontSize": "3号", "punctuation": True},
                    "level4": {"fontFamily": "仿宋", "fontSize": "3号", "punctuation": True},
                },
                "references": {"citationOrder": "先引标题后引发号", "yearBrackets": "〔 〕"},
                "attachments": {
                    "spacingBefore": "1行",
                    "indent": "左空两格",
                    "itemSuffixPunctuation": "不加标点",
                    "wrapAlign": "与文字对齐",
                    "useBookTitleMarks": "不加书名号",
                },
                "signature": {"spacingBeforeLines": "2行"},
            },
            "assistantReply": "已根据指令更新模板规则。",
            "summary": "模板规则已更新",
            "model": "deepseek-chat",
            "usage": {"total_tokens": 128},
        }

        company_resp = self.client.post("/api/units", json={"name": f"topic_rule_norm_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "规则归一化题材", "description": "测试自然语言补丁归一化"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        revised = self.client.post(
            f"/api/topics/{topic_id}/agent/revise",
            json={
                "instruction": (
                    "标题使用方正小标宋简体2号并居中；正文统一使用3号仿宋_GB2312字体，"
                    "行间距固定值28磅，首行左空两格。"
                ),
                "useDeepSeek": True,
            },
        )
        self.assertEqual(revised.status_code, 200)
        revised_rules = revised.json()["inferredRules"]

        self.assertEqual(revised_rules["title"]["fontFamily"], "方正小标宋简")
        self.assertEqual(revised_rules["title"]["fontSizePt"], 22)
        self.assertEqual(revised_rules["title"]["textAlign"], "center")
        self.assertEqual(revised_rules["title"]["arrangement"], "trapezoid")
        self.assertEqual(revised_rules["body"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(revised_rules["body"]["fontSizePt"], 16)
        self.assertEqual(revised_rules["body"]["lineSpacingPt"], 28)
        self.assertEqual(revised_rules["body"]["firstLineIndentChars"], 2)
        self.assertEqual(revised_rules["headings"]["level2"]["fontFamily"], "楷体_GB2312")
        self.assertEqual(revised_rules["headings"]["level3"]["fontSizePt"], 16)
        self.assertEqual(revised_rules["references"]["citationOrder"], "titleThenDocNo")
        self.assertEqual(revised_rules["references"]["yearBrackets"], "〔〕")
        self.assertEqual(revised_rules["attachments"]["spacingBeforeLines"], 1)
        self.assertEqual(revised_rules["attachments"]["indentChars"], 2)
        self.assertEqual(revised_rules["attachments"]["itemSuffixPunctuation"], "none")
        self.assertEqual(revised_rules["attachments"]["wrapAlign"], "text")
        self.assertFalse(revised_rules["attachments"]["useBookTitleMarks"])
        self.assertEqual(revised_rules["signature"]["spacingBeforeLines"], 2)
        self.assertNotIn("fontSize", revised_rules["title"])
        self.assertNotIn("lineHeight", revised_rules["body"])
        self.assertNotIn("textIndent", revised_rules["body"])

    def test_create_doc_from_topic_prefills_body_with_fixed_template_blocks(self) -> None:
        company_resp = self.client.post("/api/units", json={"name": f"topic_body_scaffold_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "纪要模板固定区块测试", "description": "测试固定区块自动带出"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        with SessionLocal() as session:
            template = TopicTemplate(
                topic_id=topic_id,
                version=1,
                rules={
                    "body": {"fontFamily": "仿宋_GB2312", "fontSizePt": 16},
                    "contentTemplate": {
                        "leadingNodes": [
                            {
                                "type": "paragraph",
                                "attrs": {"firstLineIndentChars": 2},
                                "content": [{"type": "text", "text": "华能云成数字产融科技（雄安）有限公司"}],
                            }
                        ],
                        "trailingNodes": [
                            {
                                "type": "paragraph",
                                "attrs": {"fontFamily": "黑体", "bold": True, "firstLineIndentChars": 2},
                                "content": [{"type": "text", "text": "主 持：金刚善"}],
                            },
                            {
                                "type": "paragraph",
                                "attrs": {"dividerRed": True},
                                "content": [],
                            },
                            {
                                "type": "paragraph",
                                "attrs": {"firstLineIndentChars": 2},
                                "content": [{"type": "text", "text": "发送：全体员工。"}],
                            }
                        ],
                        "bodyPlaceholder": "（请在此输入正文）",
                    },
                },
                source_draft_id=None,
                effective=True,
            )
            session.add(template)
            session.commit()
            session.refresh(template)
            template_id = template.id

        create_doc = self.client.post(
            f"/api/topics/{topic_id}/docs",
            json={"title": "固定区块正文", "topicTemplateId": template_id},
        )
        self.assertEqual(create_doc.status_code, 200)
        doc_id = create_doc.json()["id"]

        fetched = self.client.get(f"/api/docs/{doc_id}")
        self.assertEqual(fetched.status_code, 200)
        body = fetched.json()["body"]
        self.assertEqual(body["type"], "doc")
        self.assertEqual(len(body["content"]), 5)

        first_text = body["content"][0]["content"][0]["text"]
        placeholder_text = body["content"][1]["content"][0]["text"]
        host_text = body["content"][2]["content"][0]["text"]
        host_attrs = body["content"][2]["attrs"]
        divider_attrs = body["content"][3]["attrs"]
        last_text = body["content"][4]["content"][0]["text"]
        self.assertEqual(first_text, "华能云成数字产融科技（雄安）有限公司")
        self.assertEqual(placeholder_text, "（请在此输入正文）")
        self.assertEqual(host_text, "主 持：金刚善")
        self.assertEqual(host_attrs.get("fontFamily"), "仿宋_GB2312")
        self.assertEqual(host_attrs.get("bold"), False)
        self.assertTrue(divider_attrs.get("dividerRed"))
        self.assertEqual(last_text, "发送：全体员工。")

    def test_create_doc_from_topic_strips_title_like_leading_nodes_for_non_fixed_templates(self) -> None:
        company_resp = self.client.post("/api/units", json={"name": f"topic_strip_title_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "资源协同报告", "description": "测试动态标题剔除"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        with SessionLocal() as session:
            template = TopicTemplate(
                topic_id=topic_id,
                version=1,
                rules={
                    "title": {"fontFamily": "方正小标宋简体"},
                    "body": {"fontFamily": "仿宋_GB2312", "fontSizePt": 16},
                    "contentTemplate": {
                        "leadingNodes": [
                            {
                                "type": "paragraph",
                                "attrs": {"fontFamily": "方正小标宋简体", "textAlign": "center"},
                                "content": [{"type": "text", "text": "华能云成数字产融科技（雄安）有限公司"}],
                            },
                            {
                                "type": "paragraph",
                                "attrs": {"fontFamily": "方正小标宋简体", "textAlign": "center"},
                                "content": [{"type": "text", "text": "云成数科2025年资源协同报告"}],
                            },
                            {
                                "type": "paragraph",
                                "attrs": {"fontFamily": "黑体", "textAlign": "center"},
                                "content": [{"type": "text", "text": "2025年第1期"}],
                            },
                        ],
                        "trailingNodes": [],
                        "bodyPlaceholder": "（请在此输入正文）",
                    },
                },
                source_draft_id=None,
                effective=True,
            )
            session.add(template)
            session.commit()
            session.refresh(template)
            template_id = template.id

        create_doc = self.client.post(
            f"/api/topics/{topic_id}/docs",
            json={"topicTemplateId": template_id},
        )
        self.assertEqual(create_doc.status_code, 200)
        doc_id = create_doc.json()["id"]

        fetched = self.client.get(f"/api/docs/{doc_id}")
        self.assertEqual(fetched.status_code, 200)
        body = fetched.json()["body"]
        body_texts = [
            "".join(part.get("text", "") for part in (node.get("content") or []))
            for node in body.get("content", [])
            if isinstance(node, dict)
        ]
        self.assertIn("华能云成数字产融科技（雄安）有限公司", body_texts)
        self.assertIn("2025年第1期", body_texts)
        self.assertIn("（请在此输入正文）", body_texts)
        self.assertNotIn("云成数科2025年资源协同报告", body_texts)

    def test_delete_topic_removes_related_rows(self) -> None:
        company_resp = self.client.post("/api/units", json={"name": f"topic_delete_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "待删除题材", "description": "删除测试"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        analyze = self.client.post(
            f"/api/topics/{topic_id}/analyze",
            files=[
                (
                    "files",
                    (
                        "delete-topic-sample.docx",
                        _docx_bytes("删除题材测试样本"),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ),
                ),
            ],
        )
        self.assertEqual(analyze.status_code, 200)

        confirm = self.client.post(f"/api/topics/{topic_id}/confirm-template")
        self.assertEqual(confirm.status_code, 200)

        deleted = self.client.delete(f"/api/topics/{topic_id}")
        self.assertEqual(deleted.status_code, 200)
        self.assertEqual(deleted.json()["message"], "ok")

        fetch_deleted = self.client.get(f"/api/topics/{topic_id}")
        self.assertEqual(fetch_deleted.status_code, 404)

        with SessionLocal() as session:
            self.assertIsNone(session.query(Topic).filter(Topic.id == topic_id).first())
            self.assertEqual(session.query(TopicTemplate).filter(TopicTemplate.topic_id == topic_id).count(), 0)
            self.assertEqual(session.query(TopicTemplateDraft).filter(TopicTemplateDraft.topic_id == topic_id).count(), 0)
            self.assertEqual(session.query(DeletionAuditEvent).filter(DeletionAuditEvent.topic_id == topic_id).count(), 0)

    def test_delete_topic_template_promotes_latest_remaining_effective(self) -> None:
        company_resp = self.client.post("/api/units", json={"name": f"topic_template_delete_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "模板删除测试题材", "description": "测试模板删除"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        with SessionLocal() as session:
            t1 = TopicTemplate(
                topic_id=topic_id,
                version=1,
                rules={"body": {"fontFamily": "仿宋_GB2312", "fontSizePt": 16}},
                source_draft_id=None,
                effective=False,
            )
            t2 = TopicTemplate(
                topic_id=topic_id,
                version=2,
                rules={"body": {"fontFamily": "宋体", "fontSizePt": 16}},
                source_draft_id=None,
                effective=True,
            )
            session.add_all([t1, t2])
            session.commit()
            session.refresh(t1)
            session.refresh(t2)
            t1_id = t1.id
            t2_id = t2.id

        delete_latest = self.client.delete(f"/api/topics/{topic_id}/templates/{t2_id}")
        self.assertEqual(delete_latest.status_code, 200)
        self.assertEqual(delete_latest.json()["message"], "ok")

        list_after_first_delete = self.client.get(f"/api/topics/{topic_id}/templates")
        self.assertEqual(list_after_first_delete.status_code, 200)
        templates_after_first_delete = list_after_first_delete.json()
        self.assertEqual(len(templates_after_first_delete), 1)
        self.assertEqual(templates_after_first_delete[0]["id"], t1_id)
        self.assertTrue(templates_after_first_delete[0]["effective"])

        delete_last = self.client.delete(f"/api/topics/{topic_id}/templates/{t1_id}")
        self.assertEqual(delete_last.status_code, 200)
        self.assertEqual(delete_last.json()["message"], "ok")

        list_after_second_delete = self.client.get(f"/api/topics/{topic_id}/templates")
        self.assertEqual(list_after_second_delete.status_code, 200)
        self.assertEqual(list_after_second_delete.json(), [])

    def test_list_docs_can_filter_by_topic_id(self) -> None:
        company_resp = self.client.post("/api/units", json={"name": f"topic_docs_filter_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        topic_a = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "题材A", "description": "文档筛选A"},
        )
        self.assertEqual(topic_a.status_code, 200)
        topic_a_id = topic_a.json()["id"]

        topic_b = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "题材B", "description": "文档筛选B"},
        )
        self.assertEqual(topic_b.status_code, 200)
        topic_b_id = topic_b.json()["id"]

        create_a = self.client.post(
            f"/api/topics/{topic_a_id}/docs",
            json={"title": "题材A文档"},
        )
        self.assertEqual(create_a.status_code, 200)
        doc_a_id = create_a.json()["id"]

        create_b = self.client.post(
            f"/api/topics/{topic_b_id}/docs",
            json={"title": "题材B文档"},
        )
        self.assertEqual(create_b.status_code, 200)
        doc_b_id = create_b.json()["id"]

        filtered_a = self.client.get("/api/docs", params={"topicId": topic_a_id})
        self.assertEqual(filtered_a.status_code, 200)
        filtered_a_ids = [item["id"] for item in filtered_a.json()]
        self.assertIn(doc_a_id, filtered_a_ids)
        self.assertNotIn(doc_b_id, filtered_a_ids)

        filtered_b = self.client.get("/api/docs", params={"topicId": topic_b_id})
        self.assertEqual(filtered_b.status_code, 200)
        filtered_b_ids = [item["id"] for item in filtered_b.json()]
        self.assertIn(doc_b_id, filtered_b_ids)
        self.assertNotIn(doc_a_id, filtered_b_ids)

    @patch("app.routers.topics.extract_pdf_features", create=True)
    def test_topic_analyze_accepts_pdf_file(self, mock_extract_pdf) -> None:
        mock_extract_pdf.return_value = {
            "body": {"fontFamily": "仿宋_GB2312", "fontSizePt": 16},
            "headings": {},
            "page": {"marginsCm": {}},
        }

        company_resp = self.client.post("/api/units", json={"name": f"topic_pdf_company_{uuid.uuid4().hex[:8]}"})
        self.assertEqual(company_resp.status_code, 200)
        company_id = company_resp.json()["id"]

        created = self.client.post(
            "/api/topics",
            json={"companyId": company_id, "name": "PDF训练题材", "description": "测试PDF"},
        )
        self.assertEqual(created.status_code, 200)
        topic_id = created.json()["id"]

        analyze = self.client.post(
            f"/api/topics/{topic_id}/analyze",
            files=[
                (
                    "files",
                    (
                        "sample.pdf",
                        _fake_pdf_bytes(),
                        "application/pdf",
                    ),
                ),
            ],
        )
        self.assertEqual(analyze.status_code, 200)
        body = analyze.json()
        self.assertIn("draft", body)
        self.assertEqual(body["draft"]["status"], "draft")


if __name__ == "__main__":
    unittest.main()
