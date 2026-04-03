from __future__ import annotations

import io
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".txt"}
INLINE_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
WHOLE_BOLD_LINE_RE = re.compile(r"^\*\*(.+?)\*\*$")
HASH_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s+")
BULLET_LINE_RE = re.compile(r"^\s*[-*•]\s+")
NUMBERED_LINE_RE = re.compile(r"^\s*(?:\d+[.)、]|[（(]?[一二三四五六七八九十]+[)）]|[一二三四五六七八九十]+、)\s+")


def _normalize_text(raw_text: str) -> str:
    text = (raw_text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.strip() for line in text.split("\n")]

    normalized: list[str] = []
    last_blank = False
    for line in lines:
        if line:
            normalized.append(line)
            last_blank = False
            continue
        if not last_blank:
            normalized.append("")
        last_blank = True

    return "\n".join(normalized).strip()


def _extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))

    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        value = (paragraph.text or "").strip()
        if value:
            chunks.append(value)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                if cell_text:
                    cells.append(cell_text)
            if cells:
                chunks.append(" | ".join(cells))

    return "\n".join(chunks)


def _extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    chunks: list[str] = []
    for page in reader.pages:
        value = (page.extract_text() or "").strip()
        if value:
            chunks.append(value)
    return "\n".join(chunks)


def _extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="ignore")


def _safe_float(value: Any, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _safe_float_or_none(value: Any) -> float | None:
    try:
        if value is None:
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _normalize_color_hex(value: Any) -> str | None:
    if value is None:
        return None
    raw = str(value).strip()
    if not raw:
        return None
    if raw.startswith("#"):
        raw = raw[1:]
    if len(raw) != 6:
        return None
    return raw.upper()


def _set_run_font(run, family: str, size_pt: float, bold: bool = False, color_hex: str | None = None) -> None:
    run.font.name = family
    run._element.rPr.rFonts.set(qn("w:eastAsia"), family)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)


def _resolve_indent_pt(style: dict[str, Any], font_size_pt: float, default_indent_pt: float) -> float:
    indent_pt = _safe_float_or_none(style.get("firstLineIndentPt"))
    if indent_pt is not None:
        return indent_pt
    indent_chars = _safe_float_or_none(style.get("firstLineIndentChars"))
    if indent_chars is not None:
        return indent_chars * font_size_pt
    return default_indent_pt


def _apply_alignment(paragraph, align_value: Any, default_align: str) -> None:
    align = str(align_value or default_align).strip().lower()
    if align == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _apply_paragraph_style(paragraph, style: dict[str, Any], font_size_pt: float, default_align: str, default_indent_pt: float) -> None:
    paragraph.paragraph_format.line_spacing = Pt(_safe_float(style.get("lineSpacingPt"), 28))
    paragraph.paragraph_format.space_before = Pt(_safe_float(style.get("spaceBeforePt"), 0))
    paragraph.paragraph_format.space_after = Pt(_safe_float(style.get("spaceAfterPt"), 0))
    paragraph.paragraph_format.first_line_indent = Pt(_resolve_indent_pt(style, font_size_pt, default_indent_pt))
    _apply_alignment(paragraph, style.get("textAlign"), default_align)


def _clean_inline_markdown(text: str) -> str:
    normalized = (text or "").replace("**", "").replace("__", "").replace("`", "")
    return normalized if normalized.strip() else ""


def _parse_inline_runs(text: str) -> list[tuple[str, bool]]:
    normalized = (text or "").replace("\u00A0", " ")
    segments: list[tuple[str, bool]] = []
    cursor = 0
    for match in INLINE_BOLD_RE.finditer(normalized):
        leading = _clean_inline_markdown(normalized[cursor:match.start()])
        if leading:
            segments.append((leading, False))
        bold_text = _clean_inline_markdown(match.group(1))
        if bold_text:
            segments.append((bold_text, True))
        cursor = match.end()

    trailing = _clean_inline_markdown(normalized[cursor:])
    if trailing:
        segments.append((trailing, False))

    if segments:
        return segments

    fallback = _clean_inline_markdown(normalized)
    return [(fallback, False)] if fallback else []


def _append_runs(paragraph, segments: list[tuple[str, bool]], family: str, size_pt: float, bold: bool, color_hex: str | None = None) -> None:
    for text, segment_bold in segments:
        run = paragraph.add_run(text)
        _set_run_font(run, family, size_pt, bold=bold or segment_bold, color_hex=color_hex)


def _classify_summary_line(line: str) -> tuple[str, list[tuple[str, bool]]] | None:
    value = (line or "").strip()
    if not value:
        return None

    if HASH_HEADING_RE.match(value):
        return "heading", _parse_inline_runs(HASH_HEADING_RE.sub("", value, count=1))

    if BULLET_LINE_RE.match(value):
        return "bullet", _parse_inline_runs(BULLET_LINE_RE.sub("", value, count=1))

    if NUMBERED_LINE_RE.match(value):
        return "numbered", _parse_inline_runs(NUMBERED_LINE_RE.sub("", value, count=1))

    wrapped_heading = WHOLE_BOLD_LINE_RE.fullmatch(value)
    if wrapped_heading:
        return "heading", _parse_inline_runs(wrapped_heading.group(1))

    return "paragraph", _parse_inline_runs(value)


def _apply_list_paragraph_style(paragraph, ordered: bool) -> None:
    style_name = "List Number" if ordered else "List Bullet"
    try:
        paragraph.style = style_name
    except KeyError:
        pass


def prepare_summary_source_text(source_text: str, max_chars: int = 12000) -> dict[str, Any]:
    normalized = _normalize_text(source_text)
    if not normalized:
        raise ValueError("文档内容为空，无法生成总结")

    original_chars = len(normalized)
    truncated = original_chars > max_chars
    text = normalized[:max_chars] if truncated else normalized

    return {
        "text": text,
        "truncated": truncated,
        "originalChars": original_chars,
        "usedChars": len(text),
        "fileType": "text",
    }


def extract_text_from_uploaded_file(file_name: str, file_bytes: bytes, max_chars: int = 12000) -> dict[str, Any]:
    ext = Path(file_name or "").suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError("仅支持 DOCX、PDF、TXT 文件")

    if ext == ".docx":
        raw_text = _extract_text_from_docx(file_bytes)
    elif ext == ".pdf":
        raw_text = _extract_text_from_pdf(file_bytes)
    else:
        raw_text = _extract_text_from_txt(file_bytes)

    extracted = prepare_summary_source_text(raw_text, max_chars=max_chars)
    extracted["fileType"] = ext.replace(".", "")
    return extracted


def build_summary_docx(
    title: str,
    summary_text: str,
    source_file_name: str | None = None,
    template_rules: dict[str, Any] | None = None,
    generated_at: datetime | None = None,
) -> bytes:
    doc = Document()
    actual_title = (title or "").strip() or "公文总结"
    title_style = template_rules.get("title") if isinstance(template_rules, dict) and isinstance(template_rules.get("title"), dict) else {}
    body_style = template_rules.get("body") if isinstance(template_rules, dict) and isinstance(template_rules.get("body"), dict) else {}
    heading_styles = template_rules.get("headings") if isinstance(template_rules, dict) and isinstance(template_rules.get("headings"), dict) else {}

    title_font_family = str(title_style.get("fontFamily") or "方正小标宋简")
    title_font_size = _safe_float(title_style.get("fontSizePt"), 22)
    title_bold = bool(title_style.get("bold", False))
    title_color = _normalize_color_hex(title_style.get("colorHex"))

    p_title = doc.add_paragraph()
    _apply_paragraph_style(p_title, title_style, title_font_size, default_align="center", default_indent_pt=0)
    run = p_title.add_run(actual_title)
    _set_run_font(run, title_font_family, title_font_size, bold=title_bold, color_hex=title_color)

    now = generated_at or datetime.now(UTC)
    meta_line = f"生成时间：{now.strftime('%Y-%m-%d %H:%M:%S')} UTC"
    if source_file_name:
        meta_line = f"来源文件：{source_file_name}    {meta_line}"
    meta_paragraph = doc.add_paragraph()
    _apply_paragraph_style(meta_paragraph, body_style, 10.5, default_align="left", default_indent_pt=0)
    meta_run = meta_paragraph.add_run(meta_line)
    _set_run_font(meta_run, str(body_style.get("fontFamily") or "仿宋_GB2312"), 10.5)
    doc.add_paragraph("")

    lines = summary_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    wrote = False
    for line in lines:
        parsed = _classify_summary_line(line)
        if not parsed:
            continue
        block_type, segments = parsed
        if not segments:
            continue

        paragraph = doc.add_paragraph()
        if block_type == "heading":
            level_style = heading_styles.get("level1") if isinstance(heading_styles.get("level1"), dict) else {}
            font_family = str(level_style.get("fontFamily") or "黑体")
            font_size = _safe_float(level_style.get("fontSizePt"), 16)
            font_bold = bool(level_style.get("bold", False))
            color_hex = _normalize_color_hex(level_style.get("colorHex"))
            _apply_paragraph_style(paragraph, level_style, font_size, default_align="left", default_indent_pt=0)
            _append_runs(paragraph, segments, font_family, font_size, font_bold, color_hex=color_hex)
        else:
            font_family = str(body_style.get("fontFamily") or "仿宋_GB2312")
            font_size = _safe_float(body_style.get("fontSizePt"), 16)
            font_bold = bool(body_style.get("bold", False))
            color_hex = _normalize_color_hex(body_style.get("colorHex"))
            if block_type == "bullet":
                _apply_list_paragraph_style(paragraph, ordered=False)
                _apply_paragraph_style(paragraph, body_style, font_size, default_align="left", default_indent_pt=0)
            elif block_type == "numbered":
                _apply_list_paragraph_style(paragraph, ordered=True)
                _apply_paragraph_style(paragraph, body_style, font_size, default_align="left", default_indent_pt=0)
            else:
                _apply_paragraph_style(paragraph, body_style, font_size, default_align="left", default_indent_pt=32)
            _append_runs(paragraph, segments, font_family, font_size, font_bold, color_hex=color_hex)
        wrote = True

    if not wrote:
        paragraph = doc.add_paragraph()
        _apply_paragraph_style(paragraph, body_style, _safe_float(body_style.get("fontSizePt"), 16), default_align="left", default_indent_pt=0)
        run = paragraph.add_run("（无总结内容）")
        _set_run_font(run, str(body_style.get("fontFamily") or "仿宋_GB2312"), _safe_float(body_style.get("fontSizePt"), 16))

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()
