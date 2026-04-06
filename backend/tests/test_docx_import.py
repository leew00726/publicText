import io
import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.services.docx_import import import_docx


def _set_run_font(run, family: str, size_pt: float, *, bold: bool | None = None, color_hex: str | None = None) -> None:
    run.font.name = family
    run._element.rPr.rFonts.set(qn("w:eastAsia"), family)
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex.replace("#", ""))


def _build_sample_docx() -> bytes:
    doc = Document()

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.line_spacing = Pt(30)
    title_run = title_paragraph.add_run("践行绿色金融-全国首个绿色保理标准制定及业务创新实践")
    _set_run_font(title_run, "方正小标宋简体", 22)

    heading_paragraph = doc.add_paragraph()
    heading_paragraph.paragraph_format.space_before = Pt(6)
    heading_paragraph.paragraph_format.space_after = Pt(6)
    heading_run = heading_paragraph.add_run("一、企业介绍")
    _set_run_font(heading_run, "黑体", 16)

    body_paragraph = doc.add_paragraph()
    body_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body_paragraph.paragraph_format.line_spacing = Pt(30)
    body_paragraph.paragraph_format.first_line_indent = Pt(32)
    body_paragraph.paragraph_format.left_indent = Pt(18)
    body_paragraph.paragraph_format.right_indent = Pt(12)
    body_paragraph.paragraph_format.space_before = Pt(8)
    body_paragraph.paragraph_format.space_after = Pt(10)
    spacer_run = body_paragraph.add_run("    ")
    _set_run_font(spacer_run, "方正楷体_GB2312", 16)
    body_run = body_paragraph.add_run("华能云成数字产融科技（雄安）有限公司是中国华能集团资本公司旗下的金融科技公司。")
    _set_run_font(body_run, "仿宋_GB2312", 16, color_hex="#333333")

    plain_paragraph = doc.add_paragraph()
    plain_run = plain_paragraph.add_run("这是一个没有显式首行缩进的段落。")
    _set_run_font(plain_run, "仿宋_GB2312", 15)

    payload = io.BytesIO()
    doc.save(payload)
    return payload.getvalue()


class DocxImportTests(unittest.TestCase):
    def test_import_extracts_centered_main_title_and_preserves_paragraph_style(self) -> None:
        body, structured, report = import_docx(_build_sample_docx(), preserve_formatting=True)

        self.assertEqual(structured["title"], "践行绿色金融-全国首个绿色保理标准制定及业务创新实践")
        self.assertEqual(structured["importedTitleAttrs"]["fontFamily"], "方正小标宋简体")
        self.assertEqual(structured["importedTitleAttrs"]["fontSizePt"], 22.0)
        self.assertEqual(structured["importedTitleAttrs"]["textAlign"], "center")
        self.assertEqual(structured["importedTitleAttrs"]["lineSpacingPt"], 30.0)

        self.assertEqual(body["content"][0]["type"], "heading")
        self.assertEqual(body["content"][0]["content"][0]["text"], "一、企业介绍")
        self.assertEqual(body["content"][0]["attrs"]["level"], 1)
        self.assertEqual(body["content"][0]["attrs"]["fontFamily"], "黑体")
        self.assertEqual(body["content"][0]["attrs"]["fontSizePt"], 16.0)
        self.assertEqual(body["content"][0]["attrs"]["spaceBeforePt"], 6.0)
        self.assertEqual(body["content"][0]["attrs"]["spaceAfterPt"], 6.0)

        first_body_paragraph = body["content"][1]
        self.assertEqual(first_body_paragraph["type"], "paragraph")
        self.assertIn("华能云成数字产融科技", first_body_paragraph["content"][0]["text"])
        self.assertEqual(first_body_paragraph["attrs"]["fontFamily"], "仿宋_GB2312")
        self.assertEqual(first_body_paragraph["attrs"]["fontSizePt"], 16.0)
        self.assertEqual(first_body_paragraph["attrs"]["lineSpacingPt"], 30.0)
        self.assertEqual(first_body_paragraph["attrs"]["firstLineIndentPt"], 32.0)
        self.assertEqual(first_body_paragraph["attrs"]["leftIndentPt"], 18.0)
        self.assertEqual(first_body_paragraph["attrs"]["rightIndentPt"], 12.0)
        self.assertEqual(first_body_paragraph["attrs"]["spaceBeforePt"], 8.0)
        self.assertEqual(first_body_paragraph["attrs"]["spaceAfterPt"], 10.0)
        self.assertEqual(first_body_paragraph["attrs"]["textAlign"], "justify")
        self.assertEqual(first_body_paragraph["attrs"]["colorHex"], "#333333")

        plain_body_paragraph = body["content"][2]
        self.assertEqual(plain_body_paragraph["attrs"]["firstLineIndentPt"], 0)

        self.assertEqual(report["unrecognizedTitleCount"], 0)
        self.assertIn("已尽量保留原文段落的字体、字号、行距与缩进。", report["notes"])

    def test_import_without_preserve_formatting_keeps_lightweight_layout(self) -> None:
        body, structured, report = import_docx(_build_sample_docx(), preserve_formatting=False)

        self.assertEqual(structured["title"], "践行绿色金融-全国首个绿色保理标准制定及业务创新实践")
        self.assertIsNone(structured["importedTitleAttrs"])
        self.assertEqual(body["content"][1]["attrs"]["firstLineIndentChars"], 2)
        self.assertNotIn("fontFamily", body["content"][1]["attrs"])
        self.assertEqual(report["unrecognizedTitleCount"], 0)


if __name__ == "__main__":
    unittest.main()
