import io
import unittest

from docx import Document
from docx.oxml.ns import qn

from app.services.document_summary import build_summary_docx, extract_text_from_uploaded_file


class DocumentSummaryServiceTests(unittest.TestCase):
    @staticmethod
    def _east_asia_font(paragraph, doc: Document) -> str | None:
        run = paragraph.runs[0]
        return run._element.rPr.rFonts.get(qn("w:eastAsia"))

    @staticmethod
    def _find_paragraph(doc: Document, text: str):
        return next(paragraph for paragraph in doc.paragraphs if paragraph.text == text)

    def test_extract_text_from_txt_trims_and_truncates(self) -> None:
        payload = ("第一段。\n\n第二段。" * 2000).encode("utf-8")
        result = extract_text_from_uploaded_file("report.txt", payload, max_chars=200)

        self.assertEqual(result["fileType"], "txt")
        self.assertEqual(result["originalChars"] > 200, True)
        self.assertEqual(result["truncated"], True)
        self.assertEqual(len(result["text"]) <= 200, True)
        self.assertIn("第一段。", result["text"])

    def test_extract_text_rejects_unsupported_extension(self) -> None:
        with self.assertRaises(ValueError):
            extract_text_from_uploaded_file("report.png", b"binary")

    def test_build_summary_docx_contains_title_and_summary(self) -> None:
        raw = build_summary_docx(
            title="会议纪要总结",
            summary_text="一、核心结论\n二、后续动作",
            source_file_name="meeting.docx",
        )

        doc = Document(io.BytesIO(raw))
        texts = [p.text for p in doc.paragraphs if p.text]
        self.assertIn("会议纪要总结", texts)
        self.assertTrue(any("meeting.docx" in line for line in texts))
        self.assertTrue(any("核心结论" in line for line in texts))

    def test_build_summary_docx_renders_markdown_with_template_styles(self) -> None:
        raw = build_summary_docx(
            title="模板化总结",
            summary_text="**核心结论**\n- **平台建设**：完成升级\n- 第二项动作",
            source_file_name="memo.docx",
            template_rules={
                "title": {"fontFamily": "黑体", "fontSizePt": 24, "bold": True},
                "body": {"fontFamily": "仿宋_GB2312", "fontSizePt": 16, "lineSpacingPt": 28},
                "headings": {"level1": {"fontFamily": "楷体_GB2312", "fontSizePt": 18, "bold": True}},
            },
        )

        doc = Document(io.BytesIO(raw))
        texts = [p.text for p in doc.paragraphs if p.text]

        self.assertIn("模板化总结", texts)
        self.assertIn("核心结论", texts)
        self.assertIn("平台建设：完成升级", texts)
        self.assertIn("第二项动作", texts)
        self.assertFalse(any("**" in line for line in texts))
        self.assertFalse(any(line.strip().startswith("- ") for line in texts))

        title_paragraph = self._find_paragraph(doc, "模板化总结")
        heading_paragraph = self._find_paragraph(doc, "核心结论")
        bullet_paragraph = self._find_paragraph(doc, "平台建设：完成升级")

        self.assertEqual(self._east_asia_font(title_paragraph, doc), "黑体")
        self.assertEqual(self._east_asia_font(heading_paragraph, doc), "楷体_GB2312")
        self.assertEqual(self._east_asia_font(bullet_paragraph, doc), "仿宋_GB2312")
        self.assertEqual(bullet_paragraph.runs[0].text, "平台建设")
        self.assertTrue(bullet_paragraph.runs[0].bold)
        self.assertEqual(bullet_paragraph.runs[1].text, "：完成升级")


if __name__ == "__main__":
    unittest.main()
